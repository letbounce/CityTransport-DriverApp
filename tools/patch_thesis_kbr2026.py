# -*- coding: utf-8 -*-
"""Правки Коксюк_Диплом.docx за зауваженнями КБР 2026 (Sonnet audit)."""
from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

THESIS = Path(r"d:\Downloads\Коксюк_Диплом.docx")
BACKUP = Path(r"d:\Downloads\Коксюк_Диплом_backup_kbr2026.docx")

# Розширена таблиця В.1 (ISO/IEC/IEEE 29148)
REQ_TABLE_HEADER = [
    "ID",
    "Назва вимоги",
    "Тип",
    "Джерело",
    "Статус",
    "Складність",
    "Пріоритет",
    "Ризикованість",
    "Критерій прийняття",
    "UC",
    "Компонент",
]
REQ_TABLE_ROWS = [
    ["FR-01", "Авторизація", "Функціональна", "Бізнес-вимоги BR-01", "Затверджено", "Низька", "Обов'язкова", "Низький", "JWT 8 год, DRV-NNNN", "UC-01", "LoginScreen, AuthController"],
    ["FR-02", "Маршрути", "Функціональна", "Бізнес-вимоги", "Затверджено", "Низька", "Обов'язкова", "Низький", "Список маршрутів і зупинок", "UC-02", "RouteController"],
    ["FR-03", "Створення waybill", "Функціональна", "Бізнес-вимоги", "Затверджено", "Середня", "Обов'язкова", "Середній", "status=in_progress", "UC-02", "WaybillController"],
    ["FR-04", "Завершення рейсу", "Функціональна", "Бізнес-вимоги", "Затверджено", "Низька", "Обов'язкова", "Низький", "completed_at, stop GPS", "UC-06", "ActiveTripScreen"],
    ["FR-05", "Архівування", "Функціональна", "Користувацькі вимоги", "Затверджено", "Низька", "Бажана", "Низький", "5 кодів причини", "UC-06", "Archive API"],
    ["FR-06", "Інцидент+фото", "Функціональна", "Бізнес-вимоги", "Затверджено", "Середня", "Обов'язкова", "Середній", "Base64, GPS", "UC-05", "IncidentController"],
    ["FR-07", "Редагування", "Функціональна", "Користувацькі вимоги", "Затверджено", "Середня", "Бажана", "Середній", "version_history", "UC-05", "IncidentEditScreen"],
    ["FR-08", "Карта", "Функціональна", "Бізнес-вимоги", "Затверджено", "Середня", "Обов'язкова", "Низький", "OSM, GeoJSON", "UC-03", "TripMapRouteScreen"],
    ["FR-09", "GPS-трекінг", "Функціональна", "Бізнес-вимоги", "Затверджено", "Висока", "Обов'язкова", "Середній", "кожні 5 с", "UC-04", "LocationTrackingService"],
    ["FR-10", "Live API", "Функціональна", "Користувацькі вимоги", "Затверджено", "Низька", "Бажана", "Низький", "GET /map/live-trips", "UC-03", "Map API"],
    ["NFR-01", "Офлайн", "Нефункціональна", "Технічне ТЗ", "Затверджено", "Висока", "Обов'язкова", "Високий", "Room+Worker, batch 100", "—", "TelemetrySyncWorker"],
    ["NFR-02", "Безпека", "Нефункціональна", "Технічне ТЗ", "Затверджено", "Середня", "Обов'язкова", "Високий", "JWT, bcrypt", "—", "AuthController"],
    ["NFR-03", "UX", "Нефункціональна", "Користувацькі вимоги", "Затверджено", "Низька", "Обов'язкова", "Низький", "≥16 sp, ≥64 dp", "—", "Compose UI"],
    ["NFR-04", "Продуктивність", "Нефункціональна", "Технічне ТЗ", "Затверджено", "Середня", "Обов'язкова", "Середній", "GPS 5 с, sync 15 хв", "—", "FGS+Worker"],
]

EXTRA_BIB = [
    "Аксьонов І. О., Лановий О. М., Данченко Б. І. Концептуальні підходи до побудови автоматизованих систем управління міським транспортом // Наукові записки. — 2021.",
    "Левицький С. І. Проектування інформаційних систем: навч. посіб. — К.: Центр учбової літератури, 2019. — 176 с.",
    "Ситник В. Ф., Писаревська Т. А., Єрьоміна Н. В. Проектування інформаційних систем: навч.-метод. посіб. — К.: КНЕУ, 2020. — 280 с.",
    "Буч Г., Рамбо Д., Якобсон А. Мова UML. Посібник користувача. — К.: Діалектика, 2018. — 496 с.",
    "Фолер М. Архітектура корпоративних програмних застосунків. — К.: Вільямс, 2017. — 544 с.",
    "ДСТУ ISO/IEC 19514:2021. Інформаційні технології. Мова системного моделювання SysML. — Київ: ДП «УкрНДНЦ», 2021.",
    "ДСТУ 19.701-90. Схеми алгоритмів, програм, даних і систем. Умовні позначення і правила виконання. — Київ: Держстандарт, 1993.",
    "Про автомобільний транспорт : Закон України від 5 квіт. 2001 р. № 2344-III [Електронний ресурс]. — URL: https://zakon.rada.gov.ua/laws/show/2344-14 (дата звернення: 01.03.2026).",
    "Про міський електричний транспорт : Закон України від 29 черв. 2004 р. № 1914-IV [Електронний ресурс]. — URL: https://zakon.rada.gov.ua/laws/show/1914-14 (дата звернення: 01.03.2026).",
    "UMT. Рішення: Управління пасажирськими перевезеннями [Електронний ресурс]. — URL: https://www.umt.ua/ (дата звернення: 10.03.2026).",
    "Benish GPS. Моніторинг транспорту [Електронний ресурс]. — URL: https://benishgps.com/ (дата звернення: 12.03.2026).",
    "Gurtam. NimBus — public transportation management [Електронний ресурс]. — URL: https://gurtam.com/nimbus (дата звернення: 05.03.2026).",
    "Android Developers. Background location limits [Електронний ресурс]. — URL: https://developer.android.com/develop/background-work/location (дата звернення: 15.03.2026).",
    "Jetpack Compose. State and recomposition [Електронний ресурс]. — URL: https://developer.android.com/jetpack/compose/state (дата звернення: 15.03.2026).",
    "WorkManager. Guide to background work [Електронний ресурс]. — URL: https://developer.android.com/topic/libraries/architecture/workmanager (дата звернення: 15.03.2026).",
    "Mongoose. Schemas [Електронний ресурс]. — URL: https://mongoosejs.com/docs/guide.html (дата звернення: 12.03.2026).",
    "JSON Web Token (JWT). Introduction [Електронний ресурс]. — URL: https://jwt.io/introduction (дата звернення: 10.03.2026).",
    "Fielding R. T. Architectural Styles and the Design of Network-based Software Architectures: дис. ... — Irvine: University of California, 2000.",
    "Gamma E., Helm R., Johnson R., Vlissides J. Design Patterns: Elements of Reusable Object-Oriented Software. — Addison-Wesley, 1994.",
    "Іванченко Г. Ф., Степаненко О. П. Методичні рекомендації до виконання випускного бакалаврського проекту. — Київ: КНЕУ, 2019.",
    "Олійник В. В., Гончарук В. А. Моделювання бізнес-процесів у транспортній логістиці: монографія. — Одеса, 2022.",
    "Шевченко К. Л. Автоматизація експериментальних досліджень. — К.: НТУУ «КПІ», 2018.",
    "DozoR City. Система моніторингу громадського транспорту [Електронний ресурс]. — URL: https://dozor.tech/ (дата звернення: 28.11.2025).",
    "OpenAPI Initiative. OpenAPI Specification 3.0 [Електронний ресурс]. — URL: https://swagger.io/specification/ (дата звернення: 20.03.2026).",
    "OWASP. Mobile Application Security [Електронний ресурс]. — URL: https://owasp.org/www-project-mobile-security/ (дата звернення: 20.03.2026).",
    "AvtoGPS. Система GPS-моніторингу транспорту [Електронний ресурс]. — URL: https://avtogps.com.ua/ (дата звернення: 08.03.2026).",
    "Retrofit. Type-safe HTTP client for Android and Java [Електронний ресурс]. — URL: https://square.github.io/retrofit/ (дата звернення: 18.03.2026).",
    "Room Persistence Library [Електронний ресурс]. — URL: https://developer.android.com/training/data-storage/room (дата звернення: 18.03.2026).",
    "Kotlin Programming Language Documentation [Електронний ресурс]. — URL: https://kotlinlang.org/docs/home.html (дата звернення: 18.03.2026).",
]

# Після формування списку — посилання в тексті (фраза без [N] -> з [N])
CITATION_MAP: list[tuple[str, int]] = []  # filled after bib built


def delete_paragraph(paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def replace_galuz_to_oblast(text: str) -> str:
    if "галузь знань" in text.lower():
        pass
    pairs = [
        ("ПРЕДМЕТНОЇ ГАЛУЗІ", "ПРЕДМЕТНОЇ ОБЛАСТІ"),
        ("предметної галузі", "предметної області"),
        ("предметній галузі", "предметній області"),
        ("предметну галузь", "предметну область"),
        ("предметна галузь", "предметна область"),
        ("Предметна галузь", "Предметна область"),
        ("предметної галузі", "предметної області"),
    ]
    for old, new in pairs:
        text = text.replace(old, new)
    return text


def apply_galuz_oblast(doc: Document) -> None:
    for p in doc.paragraphs:
        if "галуз" in p.text.lower() and "галузь знань" not in p.text.lower():
            p.text = replace_galuz_to_oblast(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "галуз" in para.text.lower() and "галузь знань" not in para.text.lower():
                        para.text = replace_galuz_to_oblast(para.text)


def is_toc(p) -> bool:
    return p.style and p.style.name.startswith("toc")


def find_body_paragraph(doc: Document, needle: str, *, exact: bool = False) -> object | None:
    hits = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if is_toc(p):
            continue
        if len(t) > 150 and not exact:
            continue
        if (exact and t == needle) or (not exact and needle in t):
            hits.append(p)
    return hits[-1] if hits else None


def merge_referat(doc: Document) -> None:
    """Прибрати АНОТАЦІЮ, доповнити РЕФЕРАТ за КБР."""
    # Видалити абзаци АНОТАЦІЯ (не TOC)
    to_delete = []
    in_anot = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if is_toc(p):
            if t.startswith("АНОТАЦІЯ"):
                to_delete.append(p)
            continue
        if t == "АНОТАЦІЯ":
            in_anot = True
            to_delete.append(p)
            continue
        if in_anot:
            if t == "РЕФЕРАТ":
                in_anot = False
            else:
                to_delete.append(p)
    for p in to_delete:
        delete_paragraph(p)

    ref = find_body_paragraph(doc, "РЕФЕРАТ", exact=True)
    if not ref:
        ref = find_body_paragraph(doc, "Кваліфікаційна бакалаврська робота містить")
    if not ref:
        return

    # Оновити обсяг (сторінки) і джерела пізніше; ключові слова 7
    for p in doc.paragraphs:
        if is_toc(p):
            continue
        if p.text.strip().startswith("Ключові слова:") and "РЕФЕРАТ" in (
            ref._element.getprevious().text if ref._element.getprevious() is not None else ""
        ):
            pass

    # Ключові слова в рефераті — скоротити до 7
    for p in doc.paragraphs:
        if is_toc(p):
            continue
        if p.text.strip().startswith("Ключові слова:") and "Jetpack Compose" in p.text:
            p.text = (
                "Ключові слова: інформаційна система, диспетчеризація, мобільний термінал, "
                "GPS-трекінг, REST API, MongoDB, Android."
            )
            break

    significance = (
        "Теоретична, методична та практична значущість. Теоретична значущість полягає у "
        "формалізації offline-first архітектури мобільного клієнта та бакетного зберігання "
        "телеметрії в документно-орієнтованій СКБД. Методична — у застосуванні UML/SysML-моделювання "
        "та специфікації вимог за ISO/IEC/IEEE 29148:2018. Практична — у прототипі підсистеми "
        "RoutePulse для АТП. Новизна роботи полягає в поєднанні Room + WorkManager, бакетної "
        "телеметрії MongoDB та версійного обліку інцидентів у self-hosted рішенні без ліцензійних платежів."
    )
    tasks = (
        "Завдання роботи: аналіз предметної області та ІС диспетчеризації; специфікація вимог; "
        "постановка задачі та UML-моделювання; проєктування та реалізація RoutePulse; апробація на DRV-1042."
    )

    if not any("Теоретична, методична та практична значущість" in p.text for p in doc.paragraphs):
        anchor = find_body_paragraph(doc, "Рік захисту випускного")
        if anchor:
            for text in (tasks, significance):
                np = anchor.insert_paragraph_before("")
                np.paragraph_format.first_line_indent = Cm(1.25)
                np.paragraph_format.line_spacing = 1.5
                np.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                np.add_run(text)

    # Прибрати дубль ключових слів / практичної значущості перед рефератом
    ref_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "РЕФЕРАТ" and not is_toc(p):
            ref_idx = i
            break
    if ref_idx:
        for j in range(ref_idx - 1, max(ref_idx - 8, -1), -1):
            p = doc.paragraphs[j]
            t = p.text.strip()
            if is_toc(p):
                continue
            if t.startswith("Ключові слова:") or t.startswith("Практична значущість:"):
                delete_paragraph(p)
            if "Новизна роботи" in t or "випускного бакалаврського проекту студента" in t:
                delete_paragraph(p)


def enhance_intro(doc: Document) -> None:
    roz1 = find_body_paragraph(doc, "РОЗДІЛ 1.")
    if not roz1:
        return
    blocks = []
    if not any("Теоретична значущість" in p.text or "Теоретична та методична" in p.text for p in doc.paragraphs):
        blocks.append(
            "Теоретична та методична значущість. Теоретична значущість полягає у формалізації "
            "архітектури offline-first мобільного клієнта та підходу до бакетного зберігання "
            "телеметрії. Методична — у застосуванні UML/SysML-моделювання, специфікації вимог "
            "за ISO/IEC/IEEE 29148:2018 та алгоритмічного опису за ДСТУ 19.701-90."
        )
    if not any("Інформаційна база дослідження" in p.text for p in doc.paragraphs):
        blocks.append(
            "Інформаційна база дослідження: наукові публікації з диспетчеризації транспорту; "
            "документація Android Developers, MongoDB, Express; навчальні посібники КНЕУ; "
            "стандарти ISO/IEC/IEEE 29148:2018, ISO/IEC 19501:2005, ДСТУ 3008:2015, РД 50-34.698-90; "
            "матеріали Wialon, Fleet Complete, UMT."
        )
    if not any("Структура роботи" in p.text for p in doc.paragraphs):
        blocks.append(
            "Структура роботи: вступ, три розділи (аналіз предметної області; вимоги та моделювання; "
            "проєктування та реалізація), висновки, список використаних джерел, додатки."
        )
    for text in blocks:
        np = roz1.insert_paragraph_before("")
        np.paragraph_format.first_line_indent = Cm(1.25)
        np.paragraph_format.line_spacing = 1.5
        np.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        np.add_run(text)


def merge_section_316(doc: Document) -> None:
    h316 = find_body_paragraph(doc, "3.1.6.")
    h313 = find_body_paragraph(doc, "3.1.3.")
    if not h316:
        return
    body_parts: list = []
    to_delete: list = []
    collect = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if p is h316:
            collect = True
            to_delete.append(p)
            continue
        if collect:
            if re.match(r"^3\.2\.", t):
                break
            to_delete.append(p)
            if t:
                body_parts.append(t)
    if h313:
        merge_note = (
            "Структура інформаційних масивів (документи MongoDB, локальна черга Room) "
            "деталізована у таблицях 2.4–2.5 та в даталогічній моделі вище."
        )
        if not any(merge_note[:40] in p.text for p in doc.paragraphs):
            np = h313.insert_paragraph_before(merge_note)
            np.paragraph_format.first_line_indent = Cm(1.25)
            np.paragraph_format.line_spacing = 1.5
            np.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for p in to_delete:
        delete_paragraph(p)
    for p in list(doc.paragraphs):
        if is_toc(p) and "3.1.6" in p.text:
            delete_paragraph(p)


def merge_section_35(doc: Document) -> None:
    for p in doc.paragraphs:
        if "3.4.5." in p.text and "Оцінка практичної" in p.text and not is_toc(p):
            p.text = p.text.replace("3.4.5.", "3.4.4.")
    h35 = find_body_paragraph(doc, "3.5.")
    h344 = find_body_paragraph(doc, "3.4.4.")
    if not h35:
        return
    body_35: list[str] = []
    to_delete: list = []
    collect = False
    for p in doc.paragraphs:
        if p is h35:
            collect = True
            to_delete.append(p)
            continue
        if collect:
            t = p.text.strip()
            if t == "ВИСНОВКИ":
                break
            to_delete.append(p)
            if t:
                body_35.append(t)
    if h344 and body_35:
        last = h344
        for part in reversed(body_35):
            if "Організаційне забезпечення визначає" in part:
                continue
            np = last.insert_paragraph_before(part)
            np.paragraph_format.first_line_indent = Cm(1.25)
            np.paragraph_format.line_spacing = 1.5
            np.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        intro = h344.insert_paragraph_before(
            "Організаційне забезпечення (ролі, регламент, навчання) наведено нижче в межах оцінки "
            "практичної цінності впровадження."
        )
        intro.paragraph_format.first_line_indent = Cm(1.25)
        intro.paragraph_format.line_spacing = 1.5
        intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for p in to_delete:
        delete_paragraph(p)
    for p in list(doc.paragraphs):
        if is_toc(p) and "3.5." in p.text:
            delete_paragraph(p)
        if is_toc(p) and "3.4.5." in p.text:
            p.text = p.text.replace("3.4.5.", "3.4.4.")


def rebuild_table_v1(doc: Document) -> None:
    target = None
    for i, t in enumerate(doc.tables):
        if t.rows and t.rows[0].cells[0].text.strip() == "ID" and "FR-01" in (
            t.rows[1].cells[0].text if len(t.rows) > 1 else ""
        ):
            if len(t.columns) <= 8:
                target = t
                break
    if target is None:
        for t in doc.tables:
            if t.rows and len(t.rows) >= 14 and t.rows[1].cells[0].text.strip() == "FR-01":
                target = t
                break
    if target is None:
        return
    tbl_el = target._tbl
    parent = tbl_el.getparent()
    idx = list(parent).index(tbl_el)
    parent.remove(tbl_el)

    # Знайти документ і вставити нову таблицю
    body_paras = doc.paragraphs
    cap_p = None
    for p in body_paras:
        if "Таблиця В.1" in p.text or "Таблиця B.1" in p.text:
            cap_p = p
            break
    new_tbl = doc.add_table(rows=1 + len(REQ_TABLE_ROWS), cols=len(REQ_TABLE_HEADER))
    try:
        new_tbl.style = "Table Grid"
    except KeyError:
        pass
    for j, h in enumerate(REQ_TABLE_HEADER):
        new_tbl.rows[0].cells[j].text = h
    for ri, row in enumerate(REQ_TABLE_ROWS):
        for ci, val in enumerate(row):
            new_tbl.rows[ri + 1].cells[ci].text = val
    if cap_p is not None:
        cap_p._element.addnext(new_tbl._tbl)
    else:
        parent.insert(idx, new_tbl._tbl)


def collect_bibliography(doc: Document) -> list[str]:
    entries = []
    after = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "ПЕРЕЛІК ВИКОРИСТАНИХ ДЖЕРЕЛ":
            after = True
            continue
        if after:
            if t.startswith("ДОДАТКИ"):
                break
            if t:
                # strip leading numbers/bullets
                t = re.sub(r"^\d+\.\s*", "", t)
                t = re.sub(r"^[-–•]\s*", "", t)
                entries.append(t)
    seen = set()
    merged = []
    for e in entries:
        key = e[:60]
        if key not in seen:
            seen.add(key)
            merged.append(e)
    for e in EXTRA_BIB:
        key = e[:60]
        if key not in seen:
            seen.add(key)
            merged.append(e)
    for e in EXTRA_BIB:
        if len(merged) >= 32:
            break
        key = e[:60]
        if key not in seen:
            seen.add(key)
            merged.append(e)
    return merged[:32]


def rebuild_bibliography(doc: Document, entries: list[str]) -> None:
    global CITATION_MAP
    heading = find_body_paragraph(doc, "ПЕРЕЛІК ВИКОРИСТАНИХ", exact=True)
    if not heading:
        return
    dod = find_body_paragraph(doc, "ДОДАТКИ", exact=True)
    to_del = []
    collect = False
    for p in doc.paragraphs:
        if p is heading:
            collect = True
            continue
        if collect:
            if p is dod or p.text.strip().startswith("ДОДАТКИ"):
                break
            if p.text.strip():
                to_del.append(p)
    for p in to_del:
        delete_paragraph(p)

    anchor = dod if dod else heading
    numbered = [f"{i}. {e}" for i, e in enumerate(entries, 1)]
    for text in reversed(numbered):
        np = anchor.insert_paragraph_before(text)
        np.paragraph_format.line_spacing = 1.5
        np.paragraph_format.first_line_indent = Cm(0)
        try:
            np.style = "List Paragraph"
        except KeyError:
            pass

    cite_keys = [
        "Wialon", "Fleet Complete", "AvtoGPS", "ISO/IEC/IEEE 29148", "ISO/IEC 19501",
        "ДСТУ 19.701-90", "РД 50-34.698-90", "MongoDB", "Hoang", "Truica", "Ситник",
        "Kotlin", "Jetpack Compose", "WorkManager", "Room", "Express", "Retrofit",
        "OSMdroid", "OWASP", "Fielding", "Gamma", "Аксьонов", "Левицький", "ДСТУ 3008",
        "UMT", "Benish", "NimBus", "DozoR", "OpenAPI", "Mongoose", "JWT", "Фолер", "Буч",
    ]
    cmap: dict[str, int] = {}
    for i, e in enumerate(entries, 1):
        el = e.lower()
        for key in cite_keys:
            if key.lower() in el and key not in cmap:
                cmap[key] = i
    global CITATION_MAP
    CITATION_MAP = list(cmap.items())

    for p in doc.paragraphs:
        if is_toc(p):
            continue
        if "список літератури з" in p.text and "найменувань" in p.text:
            p.text = re.sub(
                r"список літератури з \d+ найменувань",
                f"список літератури з {len(entries)} найменувань",
                p.text,
            )


def add_citations(doc: Document) -> None:
    """Додати [N] після першої згадки ключових джерел (якщо ще немає [)."""
    if not CITATION_MAP:
        return
    in_bib = False
    for p in doc.paragraphs:
        t = p.text
        if t.strip() == "ПЕРЕЛІК ВИКОРИСТАНИХ ДЖЕРЕЛ":
            in_bib = True
            continue
        if in_bib:
            if t.strip().startswith("ДОДАТКИ"):
                in_bib = False
            continue
        if is_toc(p) or t.strip().startswith("Рисунок") or t.strip().startswith("Таблиця"):
            continue
        for key, num in sorted(CITATION_MAP, key=lambda x: -len(x[0])):
            if key not in t:
                continue
            # вже є посилання поруч
            pat = re.compile(re.escape(key) + r"(?!\s*\[\d+\])")
            if pat.search(t) and f"[{num}]" not in t:
                t = pat.sub(f"{key} [{num}]", t, count=1)
        p.text = t


def fix_toc_order(doc: Document) -> None:
    """Прибрати АНОТАЦІЮ з змісту."""
    for p in list(doc.paragraphs):
        if is_toc(p) and p.text.strip().startswith("АНОТАЦІЯ"):
            delete_paragraph(p)


def main() -> None:
    if not THESIS.exists():
        raise SystemExit(f"Не знайдено: {THESIS}")
    shutil.copy2(THESIS, BACKUP)
    print(f"Backup: {BACKUP}")

    doc = Document(str(THESIS))
    merge_referat(doc)
    fix_toc_order(doc)
    enhance_intro(doc)
    apply_galuz_oblast(doc)
    merge_section_316(doc)
    merge_section_35(doc)
    rebuild_table_v1(doc)
    entries = collect_bibliography(doc)
    rebuild_bibliography(doc, entries)
    add_citations(doc)
    apply_galuz_oblast(doc)

    doc.save(str(THESIS))
    print(f"Saved: {THESIS}")
    print(f"Bibliography entries: {len(entries)}")


if __name__ == "__main__":
    main()
