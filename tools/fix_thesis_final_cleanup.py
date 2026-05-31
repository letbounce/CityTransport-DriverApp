# -*- coding: utf-8 -*-
"""Прибрати тестування зі змісту; додати в тіло; почистити бібліографію."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

THESIS = Path(r"c:\Users\letbounce\Desktop\Коксюк_Диплом.docx")


def delete_paragraph(paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def insert_table_before(anchor, caption: str, headers, rows):
    from docx import Document as D

    p = anchor.insert_paragraph_before(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc = anchor.part.document
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        tbl.style = "Table Grid"
    except KeyError:
        pass
    for j, h in enumerate(headers):
        tbl.rows[0].cells[j].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            tbl.rows[ri + 1].cells[ci].text = val
    anchor._element.addprevious(tbl._tbl)
    anchor.insert_paragraph_before("")


def main() -> None:
    doc = Document(str(THESIS))

    # 1) Видалити зі змісту випадково вставлений текст і таблицю (після рядка змісту 3.4.2)
    in_toc = False
    to_del = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "ЗМІСТ":
            in_toc = True
            continue
        if in_toc and t.startswith("ПЕРЕЛІК УМОВНИХ"):
            in_toc = False
            break
        if in_toc and (
            t.startswith("Метою тестування")
            or t.startswith("Таблиця 3.7")
            or (t.startswith("3.4.3. Тестування") and "\t" not in t)
        ):
            to_del.append(p)
    for p in to_del:
        delete_paragraph(p)

    # TOC line with page number
    for p in doc.paragraphs:
        if "3.4.2. Результати апробації" in p.text and "\t" in p.text:
            n = p.insert_paragraph_before(
                "3.4.3. Тестування та оцінювання результатів роботи системи\t51"
            )
            n.paragraph_format.line_spacing = 1.5
            break

    # 2) Додати §3.4.3 у тіло (перед 3.4.4)
    body_has_testing = any(
        p.text.strip().startswith("Метою тестування")
        and p.style.name != "toc 1"
        for p in doc.paragraphs
    )
    for p in doc.paragraphs:
        if p.text.strip().startswith("3.4.4. Оцінка практичної") and p.style.name.startswith("Heading"):
            if not body_has_testing:
                h = p.insert_paragraph_before(
                    "3.4.3. Тестування та оцінювання результатів роботи системи"
                )
                h.style = "Heading 3"
                bp = p.insert_paragraph_before(
                    "Метою тестування є підтвердження відповідності реалізованого прототипу RoutePulse "
                    "специфікованим вимогам FR/NFR. Виконано smoke-тестування 14 API-ендпоінтів, "
                    "UX-checklist (NFR-03) та сценарії на контрольному прикладі DRV-1042 — табл. 3.7."
                )
                bp.paragraph_format.first_line_indent = Cm(1.25)
                bp.paragraph_format.line_spacing = 1.5
                insert_table_before(
                    p,
                    "Таблиця 3.7\nТест-кейси апробації підсистеми RoutePulse (DRV-1042)",
                    ["ID", "Вимога / UC", "Дія", "Очікуваний результат", "Статус"],
                    [
                        ["TC-01", "FR-01 / UC-01", "Login DRV-1042", "JWT, Home", "Пройдено"],
                        ["TC-02", "FR-03 / UC-02", "Створення waybill м.114", "HTTP 201", "Пройдено"],
                        ["TC-03", "FR-09 / UC-04", "GPS під час рейсу", "Room + sync", "Пройдено"],
                        ["TC-04", "FR-06 / UC-05", "Інцидент + фото", "HTTP 201", "Пройдено"],
                        ["TC-05", "FR-04 / UC-06", "Завершення рейсу", "completed", "Пройдено"],
                        ["TC-06", "FR-08 / UC-03", "Карта OSM", "GeoJSON", "Пройдено"],
                    ],
                )
            break

    # 3) Бібліографія: прибрати джерела перед заголовком і «чужі» рядки
    seen_heading = False
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if t == "ПЕРЕЛІК ВИКОРИСТАНИХ ДЖЕРЕЛ":
            seen_heading = True
            continue
        if not seen_heading and t and (
            "[Електронний ресурс]" in t or "— К.:" in t or t.startswith("ISO/")
        ):
            delete_paragraph(p)
            continue
        if seen_heading and t == "ДОДАТКИ":
            break
        if seen_heading and t.startswith("Розроблено специфікацію вимог"):
            delete_paragraph(p)

    doc.save(str(THESIS))
    print("Cleanup done.")


if __name__ == "__main__":
    main()
