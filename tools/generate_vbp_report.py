# -*- coding: utf-8 -*-
"""Generate bachelor's thesis report (ВБП) for RoutePulse."""
from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Mm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from vbp_extra_content import (
    EXTRA_S1,
    EXTRA_S2_ALGO,
    EXTRA_S2_REQ,
    EXTRA_S3_INFO,
    EXTRA_S3_SW,
    EXTRA_S3_TECH,
    EXTRA_S34,
)

ROOT = Path(__file__).resolve().parents[1]
DIAGRAMS = ROOT / "docs" / "diagrams"
OUT = ROOT / "docs" / "Звіт_ВБП_Коксюк_ІН-403_RoutePulse.docx"

AUTHOR = "Коксюк Олег Віталійович"
GROUP = "ІН-403"
TOPIC = (
    "Інформаційна система диспетчеризації міського пасажирського транспорту "
    "з використанням мобільних платформ"
)
PRODUCT = "RoutePulse"
PRACTICE_TXT = ROOT / "docs" / "_practice_extract.txt"


def adapt_text(t: str) -> str:
    t = t.replace("&lt;&lt;", "«").replace("&gt;&gt;", "»")
    t = t.replace("CityTransport-DriverApp", "RoutePulse")
    t = t.replace("CityDispatchDB", "CityDispatchDB (routepulse_db)")
    t = t.replace("Sigma Software", "АТП / практика Sigma Software")
    t = t.replace("Рисунок 3.1", "рис. 2.2")
    t = t.replace("Рисунок 3.2.", "рис. 2.1")
    return t


def inject_blocks(doc: Document, blocks: list[str]) -> None:
    for block in blocks:
        if block.strip():
            p(doc, adapt_text(block.strip()))


def inject_practice_range(doc: Document, start: int, end: int, skip_prefixes=()) -> None:
    if not PRACTICE_TXT.exists():
        return
    lines = PRACTICE_TXT.read_text(encoding="utf-8").splitlines()
    buf = []
    for line in lines[start - 1 : end]:
        s = line.strip()
        if not s:
            if buf:
                inject_blocks(doc, [" ".join(buf)])
                buf = []
            continue
        if s.startswith("РОЗДІЛ") or s.startswith("Рисунок") or s.startswith("Рис."):
            continue
        if any(s.startswith(x) for x in skip_prefixes):
            continue
        if s[0].isdigit() and ". " in s[:4]:
            continue
        buf.append(s)
    if buf:
        inject_blocks(doc, [" ".join(buf)])


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.space_after = Pt(0)
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        h.font.bold = True
        h.font.color.rgb = RGBColor(0, 0, 0)
        h._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        h.paragraph_format.first_line_indent = Cm(0)
        h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def p(doc: Document, text: str, *, indent: bool = True, align=None) -> None:
    para = doc.add_paragraph(text)
    if not indent:
        para.paragraph_format.first_line_indent = Cm(0)
    if align is not None:
        para.alignment = align


def p_center(doc: Document, text: str, *, bold: bool = False, size: int = 14) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Cm(0)
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def table_from_rows(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for par in hdr[i].paragraphs:
            par.runs[0].bold = True
            par.runs[0].font.name = "Times New Roman"
            par.runs[0].font.size = Pt(12)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for par in cells[ci].paragraphs:
                for r in par.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(12)
    doc.add_paragraph()


def add_figure(doc: Document, png_name: str, caption: str, width_cm: float = 16.0) -> None:
    path = DIAGRAMS / png_name
    if not path.exists():
        p(doc, f"[Рисунок відсутній: {png_name}]", indent=False)
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Cm(0)
    run = para.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    for r in cap.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
    doc.add_paragraph()


def page_break(doc: Document) -> None:
    doc.add_page_break()


def title_page(doc: Document) -> None:
    for _ in range(2):
        p_center(doc, "")
    p_center(doc, "МІНІСТЕРСТВО ОСВІТИ І НАУКИ УКРАЇНИ")
    p_center(doc, "КИЇВСЬКИЙ НАЦІОНАЛЬНИЙ ЕКОНОМІЧНИЙ УНІВЕРСИТЕТ")
    p_center(doc, "імені ВАДИМА ГЕТЬМАНА")
    p_center(doc, "Навчально-науковий інститут")
    p_center(doc, "«Інститут інформаційних технологій в економіці»")
    p_center(doc, "Кафедра інформаційних систем в економіці")
    p_center(doc, "")
    p_center(doc, "рівень вищої освіти — перший (бакалаврський)")
    p_center(doc, "галузь знань 12 «Інформаційні технології»")
    p_center(doc, "спеціальність 122 «Комп’ютерні науки»")
    p_center(doc, "")
    for _ in range(3):
        p_center(doc, "")
    p_center(doc, "ВИПУСКНИЙ БАКАЛАВРСЬКИЙ ПРОЕКТ", bold=True, size=16)
    p_center(doc, "на тему:", size=14)
    p_center(doc, f"«{TOPIC}»", bold=True, size=14)
    p_center(doc, "")
    p_center(doc, f"Програмний продукт: мобільна інформаційна підсистема «{PRODUCT}»")
    p_center(doc, "")
    for _ in range(4):
        p_center(doc, "")
    p_center(doc, f"Виконав: здобувач 4 курсу, група {GROUP}")
    p_center(doc, AUTHOR)
    p_center(doc, "")
    p_center(doc, "Керівник: _________________________________")
    p_center(doc, "(посада, прізвище, ініціали)")
    p_center(doc, "")
    p_center(doc, "Київ — 2026")


def assignment_page(doc: Document) -> None:
    page_break(doc)
    heading(doc, "ЗАВДАННЯ НА ВИПУСКНИЙ БАКАЛАВРСЬКИЙ ПРОЕКТ", 1)
    p(doc, f"Тема: «{TOPIC}».", indent=False)
    p(doc, "Об’єкт дослідження: процес диспетчеризації та оперативного управління міським пасажирським транспортом на автотранспортному підприємстві (АТП).", indent=False)
    p(doc, "Предмет дослідження: методи та засоби проєктування інформаційної підсистеми «Мобільний термінал водія» на базі Android-платформи та клієнт-серверної архітектури.", indent=False)
    p(doc, "Мета роботи: підвищення оперативності обліку рейсів, телеметрії та інцидентів шляхом розробки та апробації інформаційної підсистеми RoutePulse.", indent=False)
    p(doc, "Вихідні дані: нормативні документи з оформлення ВБП; методичні рекомендації кафедри; результати аналізу предметної галузі; технічне завдання на програмний продукт RoutePulse.", indent=False)
    p(doc, "Зміст розрахунково-пояснювальної записки: вступ; розділ 1 — характеристика предметної галузі; розділ 2 — вимоги та моделювання; розділ 3 — проєктування та реалізація; висновки; список джерел; додатки.", indent=False)
    p(doc, "Перелік графічного матеріалу: організаційна структура АТП; інформаційна модель; діаграми UML (прецеденти, класи, послідовності, діяльність); схема алгоритму; ER-моделі БД; схеми ІЗ, автоматизації, КТЗ, АРМ, мережі, структури ПЗ; діаграма трасування вимог (20 рисунків).", indent=False)
    p(doc, "Контрольний приклад: мобільний додаток RoutePulse (Android) + REST API (Node.js/Express) + MongoDB; сценарії авторизації, дорожнього листа, телеметрії, інцидентів.", indent=False)
    p(doc, "Строк подання: «___» __________ 2026 р.    Підпис студента _____________", indent=False)
    p(doc, "Завдання прийняв до виконання: «___» __________ 2026 р.    Підпис керівника _____________", indent=False)


def annotation(doc: Document) -> None:
    page_break(doc)
    heading(doc, "АНОТАЦІЯ", 1)
    p(doc, (
        f"У роботі розроблено інформаційну підсистему «{PRODUCT}» для диспетчеризації "
        "міського пасажирського транспорту. Об’єкт впровадження — диспетчерська служба АТП "
        "та мобільні робочі місця водіїв (~150 пристроїв). Реалізовано клієнт-серверну архітектуру: "
        "Android-клієнт (Kotlin, Jetpack Compose, MVVM), сервер Node.js/Express, СКБД MongoDB. "
        "Система забезпечує електронні дорожні листи, GPS-телеметрію, реєстрацію інцидентів, "
        "карту активних рейсів. Наведено моделі, схеми забезпечення та результати апробації."
    ))


def abstract(doc: Document) -> None:
    page_break(doc)
    heading(doc, "РЕФЕРАТ", 1)
    p(doc, f"Коксюк О.В. {TOPIC} : вип. бакалавр. проект. — Київ : КНЕУ, 2026. — ___ с.", indent=False)
    p(doc, (
        f"Робота присвячена проєктуванню та реалізації інформаційної підсистеми «{PRODUCT}». "
        "Мета — автоматизація обліку рейсів і оперативного інформування диспетчерської служби. "
        "Розроблено вимоги, UML-моделі, структуру БД, технічне та програмне забезпечення. "
        "Реалізовано програмний продукт і проведено контрольний приклад."
    ))
    p(doc, "Ключові слова: диспетчеризація, міський транспорт, Android, MongoDB, REST API, телеметрія, дорожній лист.", indent=False)


def table_of_contents(doc: Document) -> None:
    page_break(doc)
    heading(doc, "ЗМІСТ", 1)
    items = [
        "ВСТУП",
        "РОЗДІЛ 1. ХАРАКТЕРИСТИКА ТА АНАЛІЗ ПРЕДМЕТНОЇ ГАЛУЗІ",
        "1.1. Характеристика предметної галузі та об’єкта впровадження",
        "1.2. Аналіз існуючих рішень та обґрунтування розробки",
        "1.3. Мета, задачі дослідження та очікувані результати",
        "РОЗДІЛ 2. РОЗРОБКА ВИМОГ І МОДЕЛЮВАННЯ ІНФОРМАЦІЙНОЇ ПІДСИСТЕМИ",
        "2.1. Аналіз і специфікація вимог",
        "2.2. Постановка та алгоритм розв’язання задачі",
        "2.3. Моделювання інформаційної підсистеми",
        "РОЗДІЛ 3. ПРОЕКТУВАННЯ ТА РЕАЛІЗАЦІЯ КОМПОНЕНТІВ ПІДСИСТЕМИ",
        "3.1. Інформаційне забезпечення",
        "3.2. Технічне забезпечення",
        "3.3. Програмне забезпечення",
        "3.4. Результати реалізації та апробації",
        "ВИСНОВКИ",
        "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ",
        "ДОДАТКИ",
    ]
    for it in items:
        para = doc.add_paragraph(it)
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.tab_stops.add_tab_stop(Cm(15.0))


def abbreviations(doc: Document) -> None:
    page_break(doc)
    heading(doc, "ПЕРЕЛІК УМОВНИХ ПОЗНАЧЕНЬ, СКОРОЧЕНЬ ТА ТЕРМІНІВ", 1)
    rows = [
        ("АТП", "автотранспортне підприємство"),
        ("АРМ", "автоматизоване робоче місце"),
        ("API", "програмний інтерфейс додатків (Application Programming Interface)"),
        ("БД", "база даних"),
        ("ВБП", "випускний бакалаврський проект"),
        ("GPS", "глобальна система позиціонування"),
        ("ІЗ", "інформаційне забезпечення"),
        ("ІС", "інформаційна система"),
        ("JWT", "JSON Web Token"),
        ("КТЗ", "комплекс технічних засобів"),
        ("ПЗ", "програмне забезпечення"),
        ("REST", "архітектурний стиль передачі стану Representational State Transfer"),
        ("СКБД", "система керування базами даних"),
        ("UML", "Unified Modeling Language"),
        ("FR / NFR", "функціональна / нефункціональна вимога"),
        ("UC", "прецедент (Use Case)"),
    ]
    table_from_rows(doc, ["Скорочення", "Розшифрування"], rows)


def introduction(doc: Document) -> None:
    page_break(doc)
    heading(doc, "ВСТУП", 1)
    p(doc, (
        "Актуальність теми зумовлена потребою цифровізації процесів управління міським "
        "пасажирським транспортом у концепції «розумного міста». Паперові дорожні листи та "
        "голосовий зв’язок «водій — диспетчер» не забезпечують оперативного моніторингу руху, "
        "фіксації відхилень від графіка та швидкого реагування на інциденти. Впровадження "
        "мобільного термінала водія з передачею телеметрії та електронних повідомлень дозволяє "
        "підвищити прозорість перевезень і знизити вплив людського фактора."
    ))
    p(doc, (
        "Об’єкт дослідження — процес диспетчеризації та оперативного управління міським "
        "пасажирським транспортом на АТП."
    ))
    p(doc, (
        "Предмет дослідження — архітектурні підходи, інформаційні технології та програмні "
        "засоби побудови мобільної підсистеми «Мобільний термінал водія» RoutePulse."
    ))
    p(doc, "Мета роботи — розробити, реалізувати та апробувати інформаційну підсистему диспетчеризації з використанням мобільних платформ.")
    p(doc, "Для досягнення мети поставлено такі завдання:")
    tasks = [
        "проаналізувати предметну галузь та організаційну структуру АТП;",
        "сформулювати функціональні та нефункціональні вимоги;",
        "розробити інформаційну модель, алгоритми та UML-моделі підсистеми;",
        "спроєктувати інформаційне, технічне та програмне забезпечення;",
        "реалізувати програмний продукт RoutePulse (Android + Node.js + MongoDB);",
        "провести контрольний приклад та оцінити готовність до пілотного впровадження.",
    ]
    for i, t in enumerate(tasks, 1):
        p(doc, f"{i}. {t}")
    p(doc, (
        "Методи дослідження: аналіз предметної галузі та літературних джерел; структурно-логічне "
        "та UML-моделювання; порівняльний аналіз СКБД; математичне та алгоритмічне описування "
        "процесів; експериментальна апробація (smoke-тести, контрольний приклад)."
    ))
    p(doc, (
        "Наукова новизна полягає в комплексному узгодженні мобільного термінала водія, "
        "bucket-моделі зберігання телеметрії в MongoDB та механізму версіонування інцидентів "
        "у єдиній підсистемі для АТП."
    ))
    p(doc, (
        "Практичне значення — зменшення часу передачі даних про рейс і інцидент, можливість "
        "пілотного впровадження на парку ~150 одиниць рухомого складу."
    ))
    p(doc, (
        "Структура роботи: вступ, три розділи, висновки, список джерел, додатки та графічні матеріали."
    ))


def add_api_table(doc: Document) -> None:
    rows = [
        ("POST", "/api/auth/login", "Ні", "Авторизація, видача JWT"),
        ("GET", "/api/routes", "JWT", "Список маршрутів"),
        ("GET", "/api/routes/:id", "JWT", "Маршрут із зупинками"),
        ("GET", "/api/vehicles", "JWT", "Довідник ТЗ"),
        ("GET", "/api/waybills/active", "JWT", "Активний рейс водія"),
        ("GET", "/api/waybills/archived", "JWT", "Архів рейсів"),
        ("POST", "/api/waybills", "JWT", "Створити дорожній лист"),
        ("PATCH", "/api/waybills/:id/complete", "JWT", "Завершити рейс"),
        ("POST", "/api/waybills/:id/archive", "JWT", "Архівувати"),
        ("POST", "/api/telemetry", "JWT", "Пакет GPS (bucket)"),
        ("GET", "/api/incidents", "JWT", "Список інцидентів"),
        ("POST", "/api/incidents", "JWT", "Новий інцидент"),
        ("PATCH", "/api/incidents/:id", "JWT", "Редагування + version_history"),
        ("GET", "/api/map/live-trips", "JWT", "Маркери активних рейсів"),
        ("GET", "/health", "Ні", "Перевірка доступності API"),
    ]
    p(doc, "Повний перелік REST API RoutePulse наведено в табл. 3.2.", indent=False)
    table_from_rows(doc, ["Метод", "URL", "Auth", "Призначення"], rows)


def section1(doc: Document) -> None:
    page_break(doc)
    heading(doc, "РОЗДІЛ 1. ХАРАКТЕРИСТИКА ТА АНАЛІЗ ПРЕДМЕТНОЇ ГАЛУЗІ", 1)
    heading(doc, "1.1. Характеристика предметної галузі та об’єкта впровадження", 2)
    p(doc, (
        "Предметна галузь — міський пасажирський транспорт (автобусні перевезення). "
        "Економічний зміст задач: планування та облік рейсів, контроль дотримання маршруту "
        "і графіка, оперативне реагування на нештатні ситуації, формування звітності для "
        "керівництва АТП та диспетчерської служби."
    ))
    p(doc, (
        "Об’єкт впровадження інформаційної підсистеми RoutePulse — диспетчерська служба АТП "
        "та мобільні робочі місця водіїв. Диспетчери (~5–6 АРМ) отримують агреговані дані "
        "про активні рейси, телеметрію та інциденти; водії (~150 смартфонів) — інтерфейс "
        "керування зміною, дорожнім листом і звітністю."
    ))
    add_figure(doc, "11-org-structure.png", "Рис. 1.1. Організаційна структура АТП з виділенням диспетчерської служби")
    p(doc, (
        "На рис. 1.1 виділено диспетчерську службу як центр збору та аналізу оперативної "
        "інформації. Підрозділи експлуатації, технічного забезпечення та бухгалтерії "
        "залишаються споживачами підсумкової звітності."
    ))
    heading(doc, "1.2. Аналіз існуючих рішень та обґрунтування розробки", 2)
    p(doc, (
        "На ринку існують комплексні системи диспетчеризації (GPS-моніторинг автопарків, "
        "MIS для перевізників). Вони часто орієнтовані на диспетчерський центр і не надають "
        "водієві мінімалістичного мобільного інтерфейсу з офлайн-стійкістю. RoutePulse "
        "фокусується на терміналі водія як первинному джерелі достовірних даних."
    ))
    table_from_rows(
        doc,
        ["Рішення", "Переваги", "Недоліки для АТП"],
        [
            ("Паперовий облік", "Низька вартість", "Затримки, помилки, немає GPS"),
            ("Універсальні fleet-GPS", "Моніторинг ТЗ", "Складний UI для водія"),
            ("RoutePulse (розробка)", "Електронний наряд, інциденти, REST", "Потребує ІТ-інфраструктури"),
        ],
    )
    heading(doc, "1.3. Мета, задачі дослідження та очікувані результати", 2)
    p(doc, (
        "Мета дослідження — обґрунтувати та реалізувати інформаційну підсистему, що замінює "
        "паперовий дорожній лист і забезпечує цифровий обмін даними «водій — сервер — диспетчер»."
    ))
    p(doc, "Очікувані результати: програмний продукт RoutePulse; комплект проєктної документації; результати контрольного прикладу та smoke-тестів API.")
    p(doc, (
        "У процесі диспетчеризації формуються первинні документи: наряд водія, дорожній лист, "
        "повідомлення про відхилення від маршруту, акти інцидентів. Інформація характеризується "
        "високою динамічністю, прив’язкою до геопростору та необхідністю синхронізації між "
        "мобільними клієнтами та центром. Традиційний облік не забезпечує цілодобового "
        "моніторингу координат і часу реакції на події менше 1–2 хвилин."
    ))
    p(doc, (
        "Розподіл функцій між персоналом і технічними засобами: водій фіксує фактичні події "
        "в мобільному додатку; сервер агрегує та зберігає дані; диспетчер аналізує потоки "
        "на АРМ у LAN депо. При втраті зв’язку водій продовжує фіксацію локально з подальшою "
        "синхронізацією — це критична вимога для міських магістралей із «білими плямами» покриття."
    ))
    inject_blocks(doc, EXTRA_S1)
    inject_practice_range(doc, 162, 203)
    inject_practice_range(doc, 178, 198)


def section2_1(doc: Document) -> None:
    page_break(doc)
    heading(doc, "РОЗДІЛ 2. РОЗРОБКА ВИМОГ І МОДЕЛЮВАННЯ ІНФОРМАЦІЙНОЇ ПІДСИСТЕМИ", 1)
    heading(doc, "2.1. Аналіз і специфікація вимог", 2)
    p(doc, "Вимоги сформульовано у стилі «система повинна…» з ідентифікаторами FR/NFR.")
    table_from_rows(
        doc,
        ["ID", "Назва", "Пріоритет", "Опис"],
        [
            ("FR-01", "Авторизація", "Обов’язкова", "Вхід водія за driver_id та паролем, отримання JWT"),
            ("FR-02", "Маршрути", "Обов’язкова", "Перегляд активних маршрутів і зупинок"),
            ("FR-03", "Дорожній лист", "Обов’язкова", "Створення, активний рейс, завершення, архів"),
            ("FR-04", "Завершення рейсу", "Обов’язкова", "Фіксація completed_at, архівування"),
            ("FR-05", "Телеметрія", "Обов’язкова", "Передача GPS-пакетів під час активного рейсу"),
            ("FR-06", "Інцидент", "Обов’язкова", "Реєстрація типу, координат, опису, фото"),
            ("FR-07", "Редагування інциденту", "Бажана", "Версіонування snapshot у version_history"),
            ("FR-08", "Карта рейсів", "Бажана", "Відображення live-trips (OSMDroid)"),
            ("FR-09", "Офлайн-черга", "Обов’язкова", "Відкладена синхронізація (WorkManager)"),
            ("NFR-01", "Офлайн-стійкість", "Обов’язкова", "Кешування та повторна відправка"),
            ("NFR-02", "Безпека", "Обов’язкова", "HTTPS, JWT, захищене зберігання токена"),
            ("NFR-03", "Ергономіка", "Обов’язкова", "Великі touch-targets, мінімум відволікань"),
            ("NFR-04", "Енергоефективність", "Обов’язкова", "Оптимізація GPS-інтервалів"),
        ],
    )
    add_figure(doc, "02-use-case-diagram.png", "Рис. 2.2. Діаграма прецедентів підсистеми «Мобільний термінал водія»")
    p(doc, "Актори: Водій (первинний), Серверна підсистема (вторинний). Прецеденти узгоджені з FR-01–FR-09.")
    p(doc, (
        "Бізнес-вимоги: зменшити частку паперових дорожніх листів; забезпечити диспетчеру "
        "єдину картину активних рейсів; скоротити час реєстрації інциденту з 10–15 хв "
        "(телефонний дзвінок) до менше 1 хв (структурована форма + GPS)."
    ))
    p(doc, (
        "Користувацькі вимоги: інтерфейс українською мовою; не більше трьох кроків від входу "
        "до початку рейсу; кнопки критичних дій висотою не менше 64 dp; підтримка темного "
        "режиму для нічних змін (Material Design 3)."
    ))
    p(doc, (
        "Системні вимоги: час відгуку API < 500 мс у LAN; JWT TTL 8 год (зміна водія); "
        "обсяг БД телеметрії — до 500 тис. bucket-документів на рік при 150 ТЗ."
    ))
    inject_blocks(doc, EXTRA_S2_REQ)
    inject_practice_range(doc, 207, 234, skip_prefixes=("Об’єкт", "Предмет", "Функціональні", "Нефункціональні"))
    inject_practice_range(doc, 224, 234)


def section2_2(doc: Document) -> None:
    heading(doc, "2.2. Постановка та алгоритм розв’язання задачі", 2)
    heading(doc, "2.2.1. Постановка задачі", 3)
    p(doc, (
        "Задача: забезпечити збір, обробку та передачу оперативної інформації про виконання "
        "рейсу від мобільного термінала водія до централізованого сховища та диспетчерської служби."
    ))
    add_figure(doc, "01-information-model.png", "Рис. 2.1. Інформаційна модель розв’язання задачі")
    heading(doc, "Вихідна інформація", 3)
    table_from_rows(
        doc,
        ["Код", "Повідомлення", "Призначення", "Періодичність"],
        [
            ("D1", "Електронний дорожній лист", "Облік факту рейсу", "На подію"),
            ("D2", "Телеметрія (GPS-трек)", "Контроль маршруту", "Кожні 5–15 с"),
            ("D3", "Повідомлення про інцидент", "Реагування диспетчера", "На подію"),
            ("D4", "Архів рейсів/інцидентів", "Звітність", "За запитом"),
        ],
    )
    heading(doc, "Вхідна інформація", 3)
    table_from_rows(
        doc,
        ["Код", "Повідомлення", "Джерело", "Спосіб отримання"],
        [
            ("V1", "Довідник маршрутів", "MongoDB routes", "REST GET /api/routes"),
            ("V2", "Облікові дані водія", "MongoDB drivers", "REST POST /api/auth/login"),
            ("V3", "GPS-координати", "Модуль GNSS пристрою", "FusedLocationProvider"),
            ("V4", "Медіа інциденту", "Камера пристрою", "Локальний файл → photo_url"),
        ],
    )
    heading(doc, "2.2.2. Алгоритм розв’язання задачі", 3)
    p(doc, "Після текстового опису логіки наведено схему алгоритму обробки інциденту (ДСТУ 19.701-90).")
    add_figure(doc, "14-algorithm-incident.png", "Рис. 2.3. Схема алгоритму обробки інциденту")
    table_from_rows(
        doc,
        ["Масив", "Джерело", "Призначення"],
        [
            ("M_waybill", "waybills", "Активний/завершений рейс"),
            ("M_telemetry", "telemetry", "Bucket GPS-локацій"),
            ("M_incident", "incidents", "Журнал подій"),
        ],
    )
    inject_blocks(doc, EXTRA_S2_ALGO)
    inject_practice_range(doc, 235, 246)


def section2_3(doc: Document) -> None:
    heading(doc, "2.3. Моделювання інформаційної підсистеми", 2)
    heading(doc, "2.3.1. Моделювання поведінки системи", 3)
    add_figure(doc, "04-sequence-diagram.png", "Рис. 2.5. Діаграма послідовності «Створення дорожнього листа»")
    add_figure(doc, "05-sequence-login.png", "Рис. 2.6. Діаграма послідовності «Авторизація (Login → JWT)»")
    add_figure(doc, "06-sequence-incident-photo.png", "Рис. 2.7. Діаграма послідовності «Реєстрація інциденту з фото»")
    add_figure(doc, "07-sequence-complete-waybill.png", "Рис. 2.8. Діаграма послідовності «Завершення дорожнього листа»")
    add_figure(doc, "08-sequence-telemetry-sync.png", "Рис. 2.9. Діаграма послідовності «GPS-телеметрія та offline-черга»")
    add_figure(doc, "12-activity-parallel.png", "Рис. 2.10. Діаграма діяльності — паралельні процеси активного рейсу")
    heading(doc, "2.3.2. Моделювання структури системи", 3)
    add_figure(doc, "03-class-diagram.png", "Рис. 2.4. Діаграма класів (3-tier: Client / App Server / Data)")
    p(doc, (
        "На діаграмі класів виділено граничні класи (LoginScreen, ActiveTripScreen), "
        "класи керування (WaybillRepository, AuthViewModel) та сутності даних (Waybill, Incident, Telemetry)."
    ))
    p(doc, (
        "Архітектура клієнта: MVVM + Clean Architecture — шар presentation (Compose UI, ViewModel), "
        "domain (use cases, моделі), data (Retrofit, репозиторії, DataStore для сесії). "
        "Сервер: контролери Express, middleware requireAuth (JWT), Mongoose-моделі."
    ))
    p(doc, (
        "Сценарій «Створення дорожнього листа» (рис. 2.5): водій обирає маршрут → POST /api/waybills → "
        "запис у колекції waybills зі status=in_progress → відображення ActiveTripScreen."
    ))
    p(doc, (
        "Сценарій «Авторизація» (рис. 2.6): POST /api/auth/login → bcrypt-перевірка → видача JWT → "
        "збереження в EncryptedSharedPreferences → подальші запити з заголовком Authorization."
    ))
    p(doc, (
        "Сценарій «Телеметрія» (рис. 2.9): LocationTrackingService збирає координати → черга → "
        "WorkManager → POST /api/telemetry (bucket). На етапі апробації серверна частина повністю "
        "реалізована; клієнтський POST — у плані наступної ітерації."
    ))
    inject_practice_range(doc, 249, 270)
    heading(doc, "2.3.3. Розподіл вимог за компонентами", 3)
    add_figure(doc, "13-traceability.png", "Рис. 2.12. Діаграма трасування вимог")
    table_from_rows(
        doc,
        ["Вимога", "UC", "Компонент"],
        [
            ("FR-01", "Login", "AuthController, LoginScreen"),
            ("FR-03", "Start Trip", "WaybillController, RouteDashboard"),
            ("FR-05", "Telemetry", "TelemetryController, WorkManager"),
            ("FR-06", "Incident", "IncidentController, IncidentReportScreen"),
        ],
    )


def section3_1(doc: Document) -> None:
    page_break(doc)
    heading(doc, "РОЗДІЛ 3. ПРОЕКТУВАННЯ ТА РЕАЛІЗАЦІЯ КОМПОНЕНТІВ ПІДСИСТЕМИ", 1)
    heading(doc, "3.1. Інформаційне забезпечення", 2)
    p(doc, "Обрано СКБД MongoDB 6.x, база CityDispatchDB (логічне ім’я routepulse_db у схемах).")
    add_figure(doc, "15-info-support-scheme.png", "Рис. 3.3. Загальна схема інформаційного забезпечення")
    add_figure(doc, "09-er-infological.png", "Рис. 3.1. Інфологічна модель БД (Crow’s foot)")
    add_figure(doc, "10-er-datalogical.png", "Рис. 3.2. Даталогічна модель (колекції MongoDB)")
    p(doc, "Колекції: drivers, routes, vehicles, waybills, telemetry, incidents.")
    table_from_rows(
        doc,
        ["Масив", "Ідентифікатор", "Носій", "Макс. обсяг", "Ключі"],
        [
            ("waybills", "waybills", "MongoDB", "50 000", "driver_id, route_id, status"),
            ("telemetry", "telemetry", "MongoDB", "500 000", "waybill_id, bucket_start"),
            ("incidents", "incidents", "MongoDB", "100 000", "waybill_id, driver_id"),
        ],
    )
    table_from_rows(
        doc,
        ["Поле", "Тип", "PK/FK", "Обов’язкове", "Опис"],
        [
            ("driver_id", "String", "PK*", "так", "Табельний номер водія"),
            ("route_id", "ObjectId", "FK", "так", "Посилання на routes"),
            ("status", "enum", "—", "так", "assigned|in_progress|completed|cancelled"),
            ("locations[]", "Array", "—", "ні", "Вкладений масив GPS у telemetry"),
        ],
    )
    p(doc, "*Первинний ключ у MongoDB — _id; driver_id — бізнес-ключ з унікальним індексом.")
    p(doc, (
        "Обґрунтування MongoDB: висока швидкість append-запису телеметрії; гнучка схема "
        "для вкладених stops у routes та version_history у incidents; горизонтальне масштабування "
        "через replica set при зростанні автопарку."
    ))
    p(doc, (
        "Класифікація та кодування: driver_id (DRV-XXXX), route_number, vehicle_id (KP-XXXX), "
        "коди типів інцидентів accident|breakdown|traffic_jam|other — узгоджені з enum у Mongoose."
    ))
    p(doc, "Приклад DDL-фрагмента waybill (MongoDB/Mongoose):")
    para = doc.add_paragraph(
        '{ "driver_id": "DRV-1042", "route_id": ObjectId("..."), "status": "in_progress", '
        '"started_at": ISODate(), "vehicle_id": "KP-3204" }'
    )
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.left_indent = Cm(1.25)
    for r in para.runs:
        r.font.name = "Courier New"
        r.font.size = Pt(11)
    inject_blocks(doc, EXTRA_S3_INFO)
    inject_practice_range(doc, 271, 280)


def section3_2(doc: Document) -> None:
    heading(doc, "3.2. Технічне забезпечення", 2)
    heading(doc, "Загальні положення та схема автоматизації.", 3)
    add_figure(doc, "16-automation-scheme.png", "Рис. 3.4. Загальна схема автоматизації Android · API · MongoDB")
    heading(doc, "Структура комплексу технічних засобів.", 3)
    add_figure(doc, "17-ktz-structure.png", "Рис. 3.5. Структура КТЗ (~150 АРМ водія, 5–6 АРМ диспетчера)")
    heading(doc, "Опис автоматизованого робочого місця (АРМ) водія.", 3)
    add_figure(doc, "18-driver-workstation.png", "Рис. 3.6. Схема АРМ водія")
    table_from_rows(
        doc,
        ["Параметр", "Мінімум", "Рекомендовано"],
        [
            ("ОС", "Android 7.0 (API 24)", "Android 12+"),
            ("RAM", "2 ГБ", "4 ГБ"),
            ("Накопичувач", "32 ГБ", "64 ГБ"),
            ("Зв’язок", "4G/LTE або Wi-Fi депо", "HTTPS + JWT"),
            ("Периферія", "GPS, камера", "Кріплення в кабіні ТЗ"),
        ],
    )
    heading(doc, "Схема мережі передачі даних.", 3)
    add_figure(doc, "19-network-scheme.png", "Рис. 3.7. Схема мережі (external / DMZ / LAN)")
    p(doc, (
        "Зони: Internet (мобільні пристрої) → DMZ (API, reverse proxy, TLS) → LAN депо "
        "(MongoDB, диспетчерські АРМ). Обсяг телеметрії: орієнтовно 150×12 пак/хв у пік."
    ))
    inject_blocks(doc, EXTRA_S3_TECH)


def section3_3(doc: Document) -> None:
    heading(doc, "3.3. Програмне забезпечення", 2)
    heading(doc, "Структура програмного забезпечення.", 3)
    add_figure(doc, "20-software-structure.png", "Рис. 3.8. Структура ПЗ RoutePulse")
    heading(doc, "Системне програмне забезпечення.", 3)
    p(doc, (
        "Клієнт: Android 14 SDK, Kotlin 2.2, Jetpack Compose BOM 2024.09. "
        "Сервер: Node.js LTS, Express 4.21, Mongoose 8.6. RAD: Android Studio, MongoDB Compass, Git."
    ))
    heading(doc, "Прикладне програмне забезпечення.", 3)
    p(doc, (
        "RoutePulse App — модулі data/domain/presentation; RoutePulse API — auth, waybills, "
        "telemetry, incidents, map, vehicles. Реалізовано 18+ REST endpoints, smoke-тест server/scripts/smoke.js."
    ))
    table_from_rows(
        doc,
        ["Модуль Android", "Призначення"],
        [
            ("presentation/login", "Авторизація"),
            ("presentation/route", "Дорожні листи"),
            ("presentation/trip", "Активний рейс"),
            ("presentation/incident", "Інциденти"),
            ("presentation/map", "Карта OSMDroid"),
            ("work/TelemetrySyncWorker", "Фонова синхронізація (розширення)"),
        ],
    )
    inject_blocks(doc, EXTRA_S3_SW)
    add_api_table(doc)
    heading(doc, "3.3.4. Проєктування інтерфейсу користувача (UI/UX)", 3)
    inject_practice_range(doc, 281, 295)
    p(doc, (
        "Примітка щодо реалізації: у поточній версії RoutePulse замість Room використано DataStore "
        "для сесії; WorkManager підготовлено для телеметрії. Це відображено в розділі 3.4 як "
        "напрям доопрацювання без зміни архітектурної концепції."
    ))
    heading(doc, "3.4. Результати реалізації та апробації", 2)
    p(doc, (
        "Проведено контрольний приклад: seed БД, login, створення waybill, POST telemetry, "
        "реєстрація incident, complete waybill. Android-додаток інтегровано з API через Retrofit "
        "(base URL 10.0.2.2:3000 для емулятора)."
    ))
    p(doc, (
        "Ступінь готовності: основні FR реалізовані на сервері та в UI; повна offline-телеметрія "
        "на клієнті — у стадії доопрацювання (WorkManager-заглушка, API POST telemetry на клієнті "
        "планується). Рекомендація — пілот на 10–15 ТЗ з подальшим масштабуванням."
    ))
    inject_blocks(doc, EXTRA_S34)
    inject_practice_range(doc, 297, 336)
    heading(doc, "Організація розробки (Agile, Git)", 3)
    p(doc, (
        "Розробка RoutePulse велася ітеративно: гілки main/develop, коміти за Conventional Commits, "
        "документація в README.md та CURSOR_PROJECT_BRIEF.md. Досвід практики на Sigma Software "
        "використано для вибору MVVM, code review та Kanban-підходу до завдань."
    ))


def conclusions(doc: Document) -> None:
    page_break(doc)
    heading(doc, "ВИСНОВКИ", 1)
    items = [
        "Проаналізовано предметну галузь диспетчеризації міського транспорту та обґрунтовано доцільність мобільного термінала водія.",
        "Сформовано вимоги FR/NFR, інформаційну модель, UML-діаграми та схему алгоритму обробки інциденту.",
        "Спроєктовано інформаційне, технічне та програмне забезпечення підсистеми RoutePulse.",
        "Реалізовано програмний продукт: Android (Kotlin/Compose) + Node.js/Express + MongoDB.",
        "Проведено контрольний приклад і smoke-тести; визначено напрями доопрацювання телеметрії та офлайн-черги.",
        "Практична цінність — скорочення часу передачі інцидентів і основа для пілотного впровадження на АТП.",
    ]
    for i, t in enumerate(items, 1):
        p(doc, f"{i}. {t}")


def references(doc: Document) -> None:
    page_break(doc)
    heading(doc, "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ", 1)
    refs = [
        "ДСТУ 3008:2015. Інформація та документація. Звіти у сфері науки і техніки.",
        "ДСТУ 19.701-90. Схеми алгоритмів, програм, даних і систем.",
        "РД 50-34.698-90. Автоматизовані системи. Вимоги до змісту документів.",
        "ISO/IEC 19501:2005. Unified Modeling Language (UML).",
        "ISO/IEC/IEEE 29148:2018. Requirements engineering.",
        "Мартін Р. Чиста архітектура. Харків : Фабула, 2019. 368 с.",
        "Ситник Н.В. Проєктування баз і сховищ даних. К., 2004. 348 с.",
        "Guide to app architecture // Android Developers. URL: https://developer.android.com/topic/architecture",
        "Express – Node.js web framework. URL: https://expressjs.com/",
        "MongoDB Manual. URL: https://www.mongodb.com/docs/",
        "Методичні рекомендації до ВБП. КНЕУ, кафедра ІС в економіці, 2019.",
    ]
    for i, r in enumerate(refs, 1):
        para = doc.add_paragraph(f"{i}. {r}")
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.left_indent = Cm(1.25)
        para.paragraph_format.hanging_indent = Cm(1.25)


def appendices(doc: Document) -> None:
    page_break(doc)
    heading(doc, "ДОДАТКИ", 1)
    p(doc, "Додаток А — Лістинги ключових модулів API (server/routes, server/models).", indent=False)
    p(doc, "Додаток Б — Скріншоти екранів RoutePulse (Login, Active Trip, Incident, Map).", indent=False)
    p(doc, "Додаток В — Контрольний приклад наповнення колекцій MongoDB (seed).", indent=False)
    p(doc, "Додаток Г — Протокол інтеграційного smoke-тесту (docs/integration-smoke-test-protocol.md).", indent=False)


def build() -> Path:
    doc = Document()
    set_doc_defaults(doc)
    title_page(doc)
    assignment_page(doc)
    annotation(doc)
    abstract(doc)
    table_of_contents(doc)
    abbreviations(doc)
    introduction(doc)
    section1(doc)
    section2_1(doc)
    section2_2(doc)
    section2_3(doc)
    section3_1(doc)
    section3_2(doc)
    section3_3(doc)
    conclusions(doc)
    references(doc)
    appendices(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    import shutil
    path = build()
    desktop = Path.home() / "Desktop" / path.name
    shutil.copy2(path, desktop)
    print(f"Saved: {path}")
    print(f"Copy:  {desktop}")
