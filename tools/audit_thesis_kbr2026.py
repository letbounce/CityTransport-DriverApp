# -*- coding: utf-8 -*-
"""Audit diploma thesis against KBR 2026 methodology requirements."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

THESIS = Path(r"d:\Downloads\Коксюк_Диплом.docx")
OUT = Path(__file__).resolve().parents[1] / "docs" / "_thesis_audit_2026.txt"


def main() -> int:
    if not THESIS.exists():
        print(f"Not found: {THESIS}", file=sys.stderr)
        return 1

    doc = Document(str(THESIS))
    paras = [(p.text.strip(), p.style.name if p.style else "") for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(t for t, _ in paras)
    lines: list[str] = []

    def w(s: str = "") -> None:
        lines.append(s)

    def find_idx(needle: str, start: int = 0) -> int | None:
        for i in range(start, len(paras)):
            if needle.lower() in paras[i][0].lower():
                return i
        return None

    w("THESIS AUDIT vs KBR 2026")
    w(f"File: {THESIS}")
    w(f"Paragraphs: {len(paras)}")

    def body_idx(needle: str) -> int | None:
        hits = [i for i, (t, _) in enumerate(paras) if needle in t and len(t) < 120]
        return hits[-1] if hits else None

    idx_referat = body_idx("РЕФЕРАТ")
    idx_vstup = body_idx("ВСТУП")
    idx_roz1 = body_idx("РОЗДІЛ 1.")
    idx_vysn = body_idx("ВИСНОВКИ")
    idx_bib = body_idx("ПЕРЕЛІК ВИКОРИСТАНИХ")
    idx_dod = body_idx("ДОДАТКИ")

    w(
        f"Index REFERAT={idx_referat}, VSTUP={idx_vstup}, ROZDIL1={idx_roz1}, "
        f"VYSNOVKY={idx_vysn}, BIB={idx_bib}, DODATKY={idx_dod}"
    )

    if idx_vstup is not None and idx_dod is not None and idx_dod > idx_vstup:
        main_text = "\n".join(paras[i][0] for i in range(idx_vstup, idx_dod))
        w(f"Main part chars (VSTUP..DODATKY): {len(main_text)} (~{len(main_text) // 1800} pages est.)")
        ok_vol = 84_000 <= len(main_text) <= 110_000
        w(f"  Volume 84k-100k chars: {'OK' if ok_vol else 'CHECK (recommended 50-60 pp)'}")

    if idx_vstup is not None and idx_roz1 is not None:
        intro = "\n".join(paras[i][0] for i in range(idx_vstup, idx_roz1))
        w(f"Intro chars: {len(intro)} (~{len(intro) // 1800} pages, max ~5 pp)")
        intro_ok = len(intro) <= 9000
        w(f"  Intro length <= ~5 pages: {'OK' if intro_ok else 'TOO LONG'}")
        intro_checks = {
            "Актуальність теми": r"актуальн",
            "Аналіз останніх досліджень": r"аналіз останніх|останніх досліджень|публікацій",
            "Мета дослідження": r"мета дослідження|мета роботи|\bмета\b",
            "Завдання дослідження": r"завдання дослідження|завдання роботи|завдання:",
            "Об'єкт дослідження": r"об[''']єкт дослідження",
            "Предмет дослідження": r"предмет дослідження",
            "Методи дослідження": r"методи дослідження",
            "Теоретична значущість": r"теоретичн.*значущ",
            "Методична значущість": r"методичн.*значущ",
            "Практична значущість": r"практичн.*значущ",
            "Інформаційна база": r"інформаційн.*баз",
            "Структура роботи": r"структура роботи",
        }
        w("\n--- INTRO (required blocks per KBR 2026) ---")
        for name, pat in intro_checks.items():
            ok = bool(re.search(pat, intro, re.I))
            w(f"  [{'+' if ok else '-'}] {name}")

    w("\n--- FRONT MATTER ---")
    front = {
        "РЕФЕРАТ": r"^РЕФЕРАТ",
        "АНОТАЦІЯ (замість реферату — не відповідає)": r"^АНОТАЦІЯ",
        "ЗМІСТ": r"^ЗМІСТ",
        "Перелік умовних скорочень": r"перелік умовних|умовних скорочень",
        "Відгук керівника": r"відгук",
        "Рецензія": r"реценз",
        "Індивідуальне завдання": r"індивідуальн.*завдання",
        "Звіт подібності (плагіат)": r"подібност|плагіат",
    }
    for name, pat in front.items():
        ok = bool(re.search(pat, text, re.I | re.M))
        w(f"  [{'+' if ok else '-'}] {name}")

    req_secs = [
        "1.1.",
        "1.2.",
        "1.3.",
        "2.1.",
        "2.2.",
        "2.2.1",
        "2.2.2",
        "2.2.3",
        "2.3.",
        "2.3.1",
        "2.3.2",
        "2.3.3",
        "3.1.",
        "3.1.1",
        "3.1.2",
        "3.1.3",
        "3.1.4",
        "3.1.5",
        "3.1.6",
        "3.2.",
        "3.2.1",
        "3.2.2",
        "3.2.3",
        "3.2.4",
        "3.3.",
        "3.3.1",
        "3.3.2",
        "3.3.3",
        "3.4.",
        "3.4.1",
        "3.4.2",
        "3.4.3",
        "3.5.",
    ]
    w("\n--- SECTIONS (methodology template) ---")
    missing_secs: list[str] = []
    for s in req_secs:
        found = any(re.match(rf"^{re.escape(s)}", p[0], re.I) for p in paras)
        if not found:
            missing_secs.append(s)
            w(f"  [-] {s}")
    if not missing_secs:
        w("  [+] All required subsections present")

    extra = [p[0][:100] for p in paras if re.match(r"^3\.4\.[45]", p[0])]
    if extra:
        w("\n--- Extra headings (beyond standard 3.4.1-3.4.3) ---")
        for e in extra:
            w(f"  * {e}")

    content = {
        "Таблиця 2.1 (вихідні повідомлення)": r"таблиц[яа]\s*2\.1",
        "Таблиця 2.2 (вхідні повідомлення)": r"таблиц[яа]\s*2\.2",
        "Таблиця 3.1 (джерела даних)": r"таблиц[яа]\s*3\.1",
        "Матриця доступу": r"матриц.*доступ",
        "DDL": r"\bDDL\b",
        "SQL/CREATE TABLE": r"\bSQL\b|CREATE TABLE",
        "Контрольний приклад": r"контрольн.*приклад|DRV-1042",
        "Діаграма прецедентів": r"прецедент",
        "Діаграма послідовності": r"послідовност",
        "Діаграма класів": r"діаграм.*клас",
        "Діаграма діяльності": r"діяльност",
        "Діаграма вимог / SysML": r"SysML|діаграм.*вимог",
        "Бізнес-вимоги": r"бізнес.?вимог",
        "FR-01": r"FR-01",
        "NFR-01": r"NFR-01",
        "Організаційне забезпечення 3.5": r"організаційн.*забезпеч",
        "Класифікація та кодування": r"класифікаці.*кодуван",
        "Первинні документи": r"первинн",
        "Машинограми": r"машинограм",
        "Smoke-тест": r"smoke",
        "Графічні матеріали": r"графічн.*матеріал",
        "Новизна у висновках": r"новизн",
    }
    w("\n--- CONTENT MARKERS ---")
    gaps: list[str] = []
    for name, pat in content.items():
        ok = bool(re.search(pat, text, re.I))
        w(f"  [{'+' if ok else '-'}] {name}")
        if not ok:
            gaps.append(name)

    w(f"\nTables in docx object: {len(doc.tables)}")
    w(f"Table captions in text: {len(re.findall(r'таблиц[яа]\\s*\\d', text, re.I))}")
    w(f"Figure captions in text: {len(re.findall(r'рисунок\\s*\\d', text, re.I))}")

    bib: list[str] = []
    if idx_bib is not None:
        for i in range(idx_bib + 1, len(paras)):
            t = paras[i][0]
            if re.match(r"^ДОДАТКИ", t, re.I):
                break
            bib.append(t)
        w(f"\nBibliography entries: {len(bib)}")
        if len(bib) < 20:
            w("  WARNING: for KBR IT thesis typically 25-35+ sources recommended")

    if idx_referat is not None:
        ref_lines = [paras[i][0] for i in range(idx_referat, min(idx_referat + 12, len(paras)))]
        ref = "\n".join(ref_lines)
        w("\n--- REFERAT (first lines) ---")
        w(ref[:1500])
        ref_checks = [
            ("об'єкт", r"об[''']єкт"),
            ("предмет", r"предмет"),
            ("мета", r"\bмета\b"),
            ("завдання", r"завдан"),
            ("теоретична значущість", r"теоретичн"),
            ("практична/методична", r"практичн|методичн"),
            ("ключові слова", r"ключов"),
        ]
        w("\nReferat elements:")
        for label, pat in ref_checks:
            w(f"  [{'+' if re.search(pat, ref, re.I) else '-'}] {label}")
        m_pages = re.search(r"(\d+)\s*сторін", ref, re.I)
        m_bib = re.search(r"(\d+)\s*найменуван", ref, re.I)
        if m_pages:
            w(f"  Pages in referat: {m_pages.group(1)}")
        if m_bib and bib:
            declared = int(m_bib.group(1))
            w(f"  Sources in referat: {declared}, actual in list: {len(bib)}")
            if declared != len(bib):
                w("  MISMATCH: update referat bibliography count")

    if idx_vysn is not None and idx_bib is not None:
        concl = "\n".join(paras[i][0] for i in range(idx_vysn, idx_bib))
        w(f"\nConclusions: {len(concl)} chars, has 'новизна': {bool(re.search(r'новизн', concl, re.I))}")

    if idx_vstup is not None and idx_roz1 is not None:
        intro_text = "\n".join(paras[i][0] for i in range(idx_vstup, idx_roz1))
        w("\n--- INTRO EXCERPT ---")
        w(intro_text[:4000])

    w("\n--- SUMMARY GAPS ---")
    if gaps:
        for g in gaps:
            w(f"  * {g}")
    else:
        w("  No major content marker gaps")

    OUT.write_text("\n".join(lines), encoding="utf-8")
    print(f"Written: {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
