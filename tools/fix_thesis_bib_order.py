# -*- coding: utf-8 -*-
"""Перенести джерела з блоку після Висновків у розділ «Перелік використаних джерел»."""
from pathlib import Path

from docx import Document
from docx.shared import Cm

THESIS = Path(r"c:\Users\letbounce\Desktop\Коксюк_Диплом.docx")


def delete_paragraph(paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def main() -> None:
    doc = Document(str(THESIS))
    moved: list[str] = []
    to_delete = []
    in_misplaced = False
    bib_heading_idx = None

    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "ПЕРЕЛІК ВИКОРИСТАНИХ ДЖЕРЕЛ":
            bib_heading_idx = i
            in_misplaced = False
            break
        if t == "ВИСНОВКИ":
            in_misplaced = True
            continue
        if in_misplaced and t and t != "ПЕРЕЛІК ВИКОРИСТАНИХ ДЖЕРЕЛ":
            if "github.com" in t:
                continue
            if "Аксьонов" in t or "— К.:" in t or "[Електронний ресурс]" in t or "ISO/" in t:
                moved.append(t)
                to_delete.append(p)

    for p in reversed(to_delete):
        delete_paragraph(p)

    # append after last bibliography entry
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip() == "ДОДАТКИ" and p.style.name.startswith("Heading"):
            anchor = p
            break
    if anchor and moved:
        for text in moved:
            np = anchor.insert_paragraph_before(text)
            np.paragraph_format.line_spacing = 1.5
            try:
                np.style = "List Paragraph"
            except KeyError:
                pass

    # rename duplicate 3.4.3
    for p in doc.paragraphs:
        if p.text.strip().startswith("3.4.3. Оцінка практичної"):
            p.text = "3.4.4. Оцінка практичної цінності та перспективи впровадження"
        if "3.4.3. Оцінка практичної" in p.text and "\t" in p.text:
            p.text = p.text.replace(
                "3.4.3. Оцінка практичної",
                "3.4.4. Оцінка практичної",
            )

    # update refereat / intro counts
    for p in doc.paragraphs:
        if "список літератури з" in p.text.lower() or "13 найменувань" in p.text:
            if "13" in p.text:
                p.text = p.text.replace("13 найменувань", "32 найменування")
                p.text = p.text.replace("з 13", "з 32")
        if "Структура роботи" in p.text and "13 найменувань" in p.text:
            p.text = p.text.replace("(13 найменувань)", "(32 найменування)")

    doc.save(str(THESIS))
    print(f"Moved {len(moved)} sources; fixed 3.4.4 heading.")


if __name__ == "__main__":
    main()
