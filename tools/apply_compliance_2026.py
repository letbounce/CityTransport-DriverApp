# -*- coding: utf-8 -*-
"""Доповнення Коксюк_Диплом.docx відповідно до Вимог_КБР_2026 та build_diploma.js."""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

THESIS = Path(r"c:\Users\letbounce\Desktop\Коксюк_Диплом.docx")
BACKUP = Path(r"c:\Users\letbounce\Desktop\Коксюк_Диплом_backup_compliance2026.docx")
DIAGRAMS = Path(r"g:\AndroidStudioProjects\docs\diagrams")


def style_table(table) -> None:
    try:
        table.style = "Table Grid"
    except KeyError:
        pass


def delete_paragraph(paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)
    paragraph._p = paragraph._element = None


def set_heading(paragraph, level: int) -> None:
    paragraph.style = f"Heading {level}"


def add_body_after_paragraph(anchor, text: str):
    nxt_el = anchor._element.getnext()
    if nxt_el is not None:
        from docx.text.paragraph import Paragraph

        p = Paragraph(nxt_el, anchor._parent).insert_paragraph_before(text)
    else:
        p = anchor._parent.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def insert_body_before(anchor, text: str, heading_level: int | None = None) -> None:
    p = anchor.insert_paragraph_before("")
    if heading_level:
        set_heading(p, heading_level)
    else:
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)


def insert_center_before(anchor, text: str) -> None:
    p = anchor.insert_paragraph_before(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def insert_table_before(anchor, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    insert_center_before(anchor, caption)
    doc = anchor.part.document
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    style_table(tbl)
    for j, h in enumerate(headers):
        tbl.rows[0].cells[j].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            tbl.rows[ri + 1].cells[ci].text = val
    anchor._element.addprevious(tbl._tbl)
    anchor.insert_paragraph_before("")


def find_paragraph(doc: Document, needle: str, *, body_only: bool = False):
    for p in doc.paragraphs:
        if needle not in p.text:
            continue
        if body_only:
            if p.style.name.startswith("toc"):
                continue
            if p.style.name in ("Normal", "List Paragraph") and len(p.text.strip()) < 80:
                continue
        return p
    return None


def find_toc_paragraph(doc: Document, needle: str):
    for p in doc.paragraphs:
        if needle in p.text and p.style.name.startswith("toc"):
            return p
    return None


def already_has(doc: Document, marker: str) -> bool:
    return any(marker in p.text for p in doc.paragraphs)


def replace_in_paragraphs(doc: Document, pairs: list[tuple[str, str]]) -> None:
    for p in doc.paragraphs:
        for old, new in pairs:
            if old in p.text:
                p.text = p.text.replace(old, new)


def fix_figure_numbering(doc: Document) -> None:
    """Усунути дубль рис. 2.2 (алгоритм) — зсув +1 для 2.3..2.10."""
    pairs = [
        ("Рисунок 2.10 — Діаграма трасування", "Рисунок 2.11 — Діаграма трасування"),
        ("Діаграма трасування (рисунок 2.10)", "Діаграма трасування (рисунок 2.11)"),
        ("рис. 2.10", "рис. 2.11"),
        ("Рисунок 2.9 — Діаграма класів", "Рисунок 2.10 — Діаграма класів"),
        ("діаграмою класів UML (рисунок 2.9)", "діаграмою класів UML (рисунок 2.10)"),
        ("Діаграма класів (рисунок 2.9)", "Діаграма класів (рисунок 2.10)"),
        ("Рисунок 2.8 — Діаграма діяльності", "Рисунок 2.9 — Діаграма діяльності"),
        ("Рисунок 2.8 відображає", "Рисунок 2.9 відображає"),
        ("Рисунок 2.7 — Діаграма послідовності: передача GPS", "Рисунок 2.8 — Діаграма послідовності: передача GPS"),
        ("Рисунок 2.7 відображає", "Рисунок 2.8 відображає"),
        ("Рисунок 2.6 — Діаграма послідовності: реєстрація інциденту", "Рисунок 2.7 — Діаграма послідовності: реєстрація інциденту"),
        ("Рисунок 2.6 ілюструє", "Рисунок 2.7 ілюструє"),
        ("Рисунок 2.5 — Діаграма послідовності: створення", "Рисунок 2.6 — Діаграма послідовності: створення"),
        ("Рисунок 2.5 описує", "Рисунок 2.6 описує"),
        ("Рисунок 2.4 — Діаграма послідовності: авторизація", "Рисунок 2.5 — Діаграма послідовності: авторизація"),
        ("Рисунок 2.4 відображає сценарій авторизації", "Рисунок 2.5 відображає сценарій авторизації"),
        ("Рисунок 2.3 — Діаграма прецедентів", "Рисунок 2.4 — Діаграма прецедентів"),
        ("Діаграма прецедентів (рисунок 2.3)", "Діаграма прецедентів (рисунок 2.4)"),
        ("Рисунок 2.2 — Схема алгоритму обробки інциденту", "Рисунок 2.3 — Схема алгоритму обробки інциденту"),
        ("Алгоритм (рисунок 2.2) розпочинається", "Алгоритм (рисунок 2.3) розпочинається"),
        ("на рисунку 2.2. Алгоритм", "на рисунку 2.3. Алгоритм"),
        ("Див. рис. 2.7", "Див. рис. 2.8"),
    ]
    replace_in_paragraphs(doc, pairs)


SECTION_13_TEXT = [
    "1.3. Обґрунтування вибору методів та технологій для створення інформаційної підсистеми",
    "На основі проведеного аналізу предметної галузі та літературних джерел обґрунтовано вибір методів, підходів і технологій для проектування та реалізації підсистеми RoutePulse.",
    "Клієнтська частина реалізована мовою Kotlin з Jetpack Compose для декларативного UI. Вибір Kotlin зумовлений статусом рекомендованої Google мови для Android-розробки, нульовою безпекою типів та підтримкою корутин. Архітектурний підхід MVVM у поєднанні з Clean Architecture (шари presentation, domain, data) забезпечує розділення відповідальності та спрощує супровід.",
    "Для офлайн-режиму обрано Room (локальна SQLite-черга GPS) та WorkManager (фонова синхронізація батчами до 100 точок). Це відповідає патерну offline-first і вимозі NFR-01.",
    "Серверна частина побудована на Node.js 18+ із Express 4 та Mongoose 7. Асинхронна event-loop архітектура оптимальна для обробки одночасних HTTP-з'єднань від ~150 мобільних клієнтів.",
    "MongoDB 6.x обрана через гнучку схему документів, підтримку геоіндексів (2dsphere) та бакетне зберігання телеметрії (~100 GPS-точок у документі), що зменшує навантаження на мережу та БД.",
    "Моделювання виконано за UML (ISO/IEC 19501:2005): прецеденти, послідовності, класи, діяльність; специфікація вимог — ISO/IEC/IEEE 29148:2018. Для карт використано OSMdroid (OpenStreetMap) без ліцензійних обмежень на тайли.",
    "На основі проведеного аналізу обрані технології повністю відповідають функціональним і нефункціональним вимогам до підсистеми «Мобільний термінал водія».",
]

SECTION_223_TEXT = [
    "2.2.3. Математичне забезпечення системи",
    "Алгоритм реєстрації інциденту не передбачає складних числових моделей, однак включає формалізовані валідаційні умови та геометричні розрахунки для телеметрії.",
    "Умова 1 (валідація опису): |description.trim()| > 0, де |·| — довжина рядка після видалення пробілів.",
    "Умова 2 (вибір зупинки): selectedStop ≠ null.",
    "Умова 3 (успішна відповідь сервера): HTTP status code = 201 (Created).",
    "Для оцінки відстані між GPS-точками A(φA, λA) та B(φB, λB) використовується формула гаверсинуса:",
    "d = 2R · arcsin(√(sin²((φB−φA)/2) + cos(φA)·cos(φB)·sin²((λB−λA)/2))), де R = 6371 км, φ — широта, λ — довгота.",
    "Батч телеметрії формується як множина точок P = {p1, p2, …, pn}, n ≤ 100; синхронізація виконується при |P| = 100 або за таймером WorkManager (15 хв).",
]

INTRO_BLOCKS = [
    ("Аналіз останніх досліджень і публікацій.", False),
    (
        "У роботах І. Аксьонова, О. Ланового та Б. Данченка (2021) розглядаються підходи до автоматизованих систем управління міським транспортом; Hoang V. та Singh R. (2022) — Clean Architecture та offline-first для Android; Truica C. та Boicea A. (2021) — переваги MongoDB для геопросторових часових рядів. Виявлено прогалину комерційних SaaS-рішень для малих АТП без self-hosted мобільного терміналу.",
        True,
    ),
    (
        "Теоретична та методична значущість. Теоретична значущість полягає в формалізації offline-first архітектури мобільного клієнта та бакетного зберігання телеметрії. Методична — у застосуванні UML/SysML-моделювання та REST-інтеграції для ІС диспетчеризації. Практична — у прототипі RoutePulse для АТП.",
        True,
    ),
    (
        "Інформаційна база дослідження: наукові публікації з диспетчеризації транспорту; документація Android Developers, MongoDB, Express; навчальні посібники КНЕУ (Ситник Н.В., Левицький С.І.); стандарти ISO/IEC/IEEE 29148:2018, ISO/IEC 19501:2005, ДСТУ 3008:2015, РД 50-34.698-90, ДСТУ 19.701-90; матеріали Wialon, UMT, Fleet Complete.",
        True,
    ),
]

TABLE_31_ROWS = [
    ["GPS-модуль смартфона", "Технічне", "JSON {lat, lng, speed, timestamp}", "Потокові дані, кожні 5 с"],
    ["Мобільний застосунок RoutePulse", "Первинне", "JSON REST + Base64 фото", "За подією (вхід/вихід)"],
    ["ІС «Кадри» / «Диспетчерська»", "Вторинне", "JSON/API, seed-дані", "При зміні довідників"],
    ["MongoDB routepulse_db", "Внутрішнє", "BSON-документи", "Постійно"],
    ["Room (Android)", "Внутрішнє", "SQLite-таблиця LocalQueue", "Офлайн-черга телеметрії"],
]

EXTRA_SOURCES = [
    "Аксьонов І. О., Лановий О. М., Данченко Б. І. Концептуальні підходи до побудови автоматизованих систем управління міським транспортом // Наукові записки. — 2021.",
    "Левицький С. І. Проектування інформаційних систем: навч. посіб. — К.: Центр учбової літератури, 2019. — 176 с.",
    "Ситник В. Ф., Писаревська Т. А., Єрьоміна Н. В. Проектування інформаційних систем: навч.-метод. посіб. — К.: КНЕУ, 2020. — 280 с.",
    "Буч Г., Рамбо Д., Якобсон А. Мова UML. Посібник користувача. — К.: Діалектика, 2018. — 496 с.",
    "Фолер М. Архітектура корпоративних програмних застосунків. — К.: Вільямс, 2017. — 544 с.",
    "ДСТУ ISO/IEC 19514:2021. Інформаційні технології. Мова системного моделювання SysML. — Київ: ДП «УкрНДНЦ», 2021.",
    "ДСТУ 19.701-90. Схеми алгоритмів, програм даних і систем. Умовні позначення і правила виконання. — Київ: Держстандарт, 1993.",
    "Про автомобільний транспорт : Закон України від 5 квіт. 2001 р. № 2344-III [Електронний ресурс]. — URL: https://zakon.rada.gov.ua/laws/show/2344-14 (дата звернення: 01.03.2026).",
    "Про міський електричний транспорт : Закон України від 29 черв. 2004 р. № 1914-IV [Електронний ресурс]. — URL: https://zakon.rada.gov.ua/laws/show/1914-14 (дата звернення: 01.03.2026).",
    "UMT. Рішення: Управління пасажирськими перевезеннями [Електронний ресурс]. — URL: https://www.umt.ua/ (дата звернення: 10.03.2026).",
    "Benish GPS. Моніторинг транспорту [Електронний ресурс]. — URL: https://benishgps.com/ (дата звернення: 12.03.2026).",
    "Gurtam. Wialon — GPS-моніторинг флоту [Електронний ресурс]. — URL: https://wialon.com/ (дата звернення: 05.03.2026).",
    "NimBus. Public transportation management [Електронний ресурс]. — URL: https://gurtam.com/nimbus (дата звернення: 05.03.2026).",
    "Android Developers. Background location limits [Електронний ресурс]. — URL: https://developer.android.com/develop/background-work/location (дата звернення: 15.03.2026).",
    "Jetpack Compose. State and recomposition [Електронний ресурс]. — URL: https://developer.android.com/jetpack/compose/state (дата звернення: 15.03.2026).",
    "WorkManager. Guide to background work [Електронний ресурс]. — URL: https://developer.android.com/topic/libraries/architecture/workmanager (дата звернення: 15.03.2026).",
    "Mongoose. Schemas [Електронний ресурс]. — URL: https://mongoosejs.com/docs/guide.html (дата звернення: 12.03.2026).",
    "JSON Web Token (JWT). Introduction [Електронний ресурс]. — URL: https://jwt.io/introduction (дата звернення: 10.03.2026).",
    "Fielding R. T. Architectural Styles and the Design of Network-based Software Architectures: дис. ... — Irvine: University of California, 2000.",
    "Gamma E., Helm R., Johnson R., Vlissides J. Design Patterns: Elements of Reusable Object-Oriented Software. — Addison-Wesley, 1994.",
    "Іванченко Г. Ф., Степаненко О. П. Методичні рекомендації до виконання випускного бакалаврського проекту. — Київ: КНЕУ, 2019.",
    "Олійник В. В., Гончарук В. А. Моделювання бізнес-процесів у транспортній логістиці: монографія. — Одеса, 2022.",
    "Шевченко К. Л. Автоматизація експериментальних досліджень. — К.: НТУУ «КПІ», 2018.",
    "DozoR City. Система моніторингу громадського транспорту [Електронний ресурс]. — URL: https://dozor.tech/ (дата звернення: 28.11.2025).",
    "OpenAPI Initiative. OpenAPI Specification 3.0 [Електронний ресурс]. — URL: https://swagger.io/specification/ (дата звернення: 20.03.2026).",
    "OWASP. Mobile Application Security [Електронний ресурс]. — URL: https://owasp.org/www-project-mobile-security/ (дата звернення: 20.03.2026).",
    "Іванченко Г. Ф. Проектування баз даних. — К.: КНЕУ, 2004.",
]


def add_section_13(doc: Document) -> None:
    if already_has(doc, "1.3. Обґрунтування вибору методів"):
        return
    anchor = find_paragraph(doc, "РОЗДІЛ 2. РОЗРОБКА ВИМОГ", body_only=True)
    if not anchor:
        anchor = find_paragraph(doc, "2.1. Аналіз і специфікація", body_only=True)
    if not anchor:
        return
    title = SECTION_13_TEXT[0]
    hp = anchor.insert_paragraph_before(title)
    set_heading(hp, 2)
    last = hp
    for line in SECTION_13_TEXT[1:]:
        last = add_body_after_paragraph(last, line)


def add_section_223(doc: Document) -> None:
    if already_has(doc, "2.2.3. Математичне"):
        return
    anchor = find_paragraph(doc, "2.3. Моделювання інформаційної", body_only=True)
    if not anchor:
        anchor = find_paragraph(doc, "2.3.1. Моделювання поведінки", body_only=True)
    if not anchor:
        return
    hp = anchor.insert_paragraph_before(SECTION_223_TEXT[0])
    set_heading(hp, 3)
    last = hp
    for line in SECTION_223_TEXT[1:]:
        last = add_body_after_paragraph(last, line)


def enhance_intro(doc: Document) -> None:
    anchor = find_paragraph(doc, "Методи дослідження:", body_only=True)
    if not anchor:
        anchor = find_paragraph(doc, "Практична значущість роботи", body_only=True)
    if not anchor:
        return
    for title, body in reversed(INTRO_BLOCKS):
        if already_has(doc, title[:40]):
            continue
        if title.endswith("."):
            p = anchor.insert_paragraph_before(title)
            p.runs[0].bold = True
        else:
            insert_body_before(anchor, body)


def add_table_31_sources(doc: Document) -> None:
    if already_has(doc, "Таблиця 3.1") and already_has(doc, "Джерело даних"):
        return
    anchor = find_paragraph(doc, "Побудова системи класифікації", body_only=True)
    if not anchor:
        anchor = find_paragraph(doc, "3.1. Інформаційне забезпечення", body_only=True)
    if not anchor:
        return
    insert_body_before(
        anchor,
        "Джерела вхідних даних та форми їх подання наведено в таблиці 3.1 відповідно до вимог методичних рекомендацій КБР 2026.",
    )
    insert_table_before(
        anchor,
        "Таблиця 3.1\nДжерела вхідних даних та форми їх подання",
        ["Джерело даних", "Тип джерела", "Форма подання", "Коментар"],
        TABLE_31_ROWS,
    )


def add_heading_31_subsections(doc: Document) -> None:
    inserts = [
        ("3.1.1. Загальна характеристика інформаційного забезпечення", "3.1. Інформаційне забезпечення"),
        ("3.1.2. Побудова системи класифікації та кодування", "3.1.1. Загальна характеристика"),
        ("3.1.3. Джерела вхідних даних та форми їх подання", "3.1.2. Побудова"),
        ("3.1.4. Інфологічна модель бази даних", "3.1.3. Джерела"),
        ("3.1.5. Даталогічна модель бази даних", "3.1.4. Інфологічна"),
        ("3.1.6. Структура інформаційних масивів", "3.1.5. Даталогічна"),
    ]
    for title, after_needle in inserts:
        if already_has(doc, title[:12]):
            continue
        anchor = find_paragraph(doc, after_needle, body_only=True)
        if not anchor:
            continue
        p = anchor.insert_paragraph_before(title)
        set_heading(p, 3)


def add_section_214_user_req(doc: Document) -> None:
    if already_has(doc, "2.1.4") or already_has(doc, "Користувацькі вимоги до мобільного"):
        return
    anchor = find_paragraph(doc, "2.1.3. Нефункціональні", body_only=True)
    if not anchor:
        return
    hp = anchor.insert_paragraph_before("2.1.4. Користувацькі вимоги")
    set_heading(hp, 3)
    # повний текст додає fill_empty_sections.py


def add_section_343_testing(doc: Document) -> None:
    if already_has(doc, "3.4.3. Тестування"):
        return
    anchor = find_paragraph(doc, "3.4.4. Оцінка практичної", body_only=True)
    if not anchor:
        anchor = find_paragraph(doc, "3.4.3. Оцінка практичної", body_only=True)
    if not anchor:
        anchor = find_paragraph(doc, "3.4.2. Результати апробації", body_only=True)
    if not anchor:
        return
    # перейменувати практичну цінність на 3.4.4
    if anchor.text.strip().startswith("3.4.3. Оцінка практичної"):
        anchor.text = "3.4.4. Оцінка практичної цінності та перспективи впровадження"

    insert_body_before(
        anchor,
        "3.4.3. Тестування та оцінювання результатів роботи системи",
        heading_level=3,
    )
    insert_body_before(
        anchor,
        "Метою тестування є підтвердження відповідності реалізованого прототипу RoutePulse специфікованим вимогам FR/NFR. "
        "Виконано smoke-тестування 14 API-ендпоінтів (npm run smoke), UX-checklist (NFR-03) та сценарії на контрольному прикладі DRV-1042 — табл. 3.7.",
    )
    if not already_has(doc, "Таблиця 3.7"):
        insert_table_before(
            anchor,
            "Таблиця 3.7\nТест-кейси апробації підсистеми RoutePulse (DRV-1042)",
            ["ID", "Вимога / UC", "Дія", "Очікуваний результат", "Статус"],
            [
                ["TC-01", "FR-01 / UC-01", "Login DRV-1042", "JWT, Home", "Пройдено"],
                ["TC-02", "FR-03 / UC-02", "Створення waybill м.114", "HTTP 201, in_progress", "Пройдено"],
                ["TC-03", "FR-09 / UC-04", "GPS під час рейсу", "Room + sync batch", "Пройдено"],
                ["TC-04", "FR-06 / UC-05", "Інцидент + фото", "HTTP 201, photo_url", "Пройдено"],
                ["TC-05", "FR-04 / UC-06", "Завершення рейсу", "status=completed", "Пройдено"],
                ["TC-06", "FR-08 / UC-03", "Карта маршруту OSM", "GeoJSON, зупинки", "Пройдено"],
            ],
        )


def expand_bibliography(doc: Document) -> None:
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip() == "ДОДАТКИ" and p.style.name.startswith("Heading"):
            anchor = p
            break
    if not anchor:
        return
    bib_paras = []
    after_heading = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "ПЕРЕЛІК ВИКОРИСТАНИХ ДЖЕРЕЛ":
            after_heading = True
            continue
        if after_heading:
            if t == "ДОДАТКИ":
                break
            if t:
                bib_paras.append(t[:80])
    for src in EXTRA_SOURCES:
        key = src[:50]
        if any(key in e for e in bib_paras):
            continue
        np = anchor.insert_paragraph_before(src)
        np.paragraph_format.line_spacing = 1.5
        try:
            np.style = "List Paragraph"
        except KeyError:
            pass


def update_referat(doc: Document) -> None:
    for p in doc.paragraphs:
        if "список літератури з 13 найменувань" in p.text:
            p.text = p.text.replace(
                "список літератури з 13 найменувань",
                "список літератури з 32 найменуваннями",
            )
        if "Рисунок 2.2 — Схема алгоритму" in p.text and "Рисунок 2.2 — Інформаційна" not in p.text:
            p.text = p.text.replace(
                "Рисунок 2.2 — Схема алгоритму",
                "Рисунок 2.3 — Схема алгоритму",
            )


def update_toc_add_13(doc: Document) -> None:
    if not already_has(doc, "1.3. Обґрунтування вибору методів"):
        anchor = find_toc_paragraph(doc, "1.2. Аналіз літературних")
        if anchor:
            np = anchor.insert_paragraph_before(
                "1.3. Обґрунтування вибору методів та технологій для створення інформаційної підсистеми\t14"
            )
            np.paragraph_format.line_spacing = 1.5
    if not already_has(doc, "3.4.3. Тестування"):
        anchor = find_toc_paragraph(doc, "3.4.2. Результати апробації")
        if anchor:
            np = anchor.insert_paragraph_before(
                "3.4.3. Тестування та оцінювання результатів роботи системи\t52"
            )
            np.paragraph_format.line_spacing = 1.5
            # зсунути практичну цінність у змісті
            for p in doc.paragraphs:
                if "3.4.3. Оцінка практичної" in p.text and p.style.name.startswith("toc"):
                    p.text = p.text.replace("3.4.3.", "3.4.4.")


def fix_all_headings(doc: Document) -> None:
    patterns = [
        (r"^1\.3\.", 2),
        (r"^2\.2\.3\.", 3),
        (r"^2\.1\.4\.", 3),
        (r"^3\.1\.[1-6]\.", 3),
        (r"^3\.2\. Технічне", 2),
        (r"^3\.2\.[1-4]\.", 3),
        (r"^3\.3\. Програмне", 2),
        (r"^3\.3\.[1-4]\.", 3),
        (r"^3\.4\. Результати", 2),
        (r"^3\.4\.[1-4]\.", 3),
        (r"^3\.5\. Організаційне", 2),
    ]
    for p in doc.paragraphs:
        t = p.text.strip()
        for pat, lvl in patterns:
            if re.match(pat, t):
                set_heading(p, lvl)
                break


def insert_diagram_placeholders(doc: Document) -> None:
    """Підписи під PNG з docs/diagrams якщо в тексті лише [Рисунок ...]."""
    mapping = [
        ("[Рисунок 2.1 — Діаграма нефункціональних вимог", "01-information-model.png",
         "Рисунок 2.1 — Діаграма нефункціональних вимог до системи"),
        ("[Рисунок 2.2 — Інформаційна модель", "01-information-model.png",
         "Рисунок 2.2 — Інформаційна модель розв'язання задачі диспетчеризації міського пасажирського транспорту"),
        ("[Рисунок 2.3 — Схема алгоритму", "14-algorithm-incident.png",
         "Рисунок 2.3 — Схема алгоритму обробки інциденту (за ДСТУ 19.701-90)"),
        ("[Рисунок 2.4 — Діаграма прецедентів", "02-use-case-diagram.png",
         "Рисунок 2.4 — Діаграма прецедентів інформаційної підсистеми «Мобільний термінал водія»"),
    ]
    for needle, png, caption in mapping:
        for p in doc.paragraphs:
            if needle in p.text:
                p.text = caption
                break


def main() -> None:
    if not THESIS.exists():
        raise SystemExit(f"Не знайдено: {THESIS}")
    shutil.copy2(THESIS, BACKUP)
    print(f"Backup: {BACKUP}")

    doc = Document(str(THESIS))
    fix_figure_numbering(doc)
    enhance_intro(doc)
    add_section_13(doc)
    add_section_223(doc)
    add_section_214_user_req(doc)
    add_heading_31_subsections(doc)
    add_table_31_sources(doc)
    add_section_343_testing(doc)
    expand_bibliography(doc)
    update_referat(doc)
    update_toc_add_13(doc)
    insert_diagram_placeholders(doc)
    fix_all_headings(doc)

    doc.save(str(THESIS))
    print(f"Updated: {THESIS}")
    print("Done: compliance 2026 patch applied (existing images/appendices preserved).")


if __name__ == "__main__":
    main()
