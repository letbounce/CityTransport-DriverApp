# -*- coding: utf-8 -*-
"""Append Section 3 to thesis docx with diagrams."""
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

THESIS = Path(r"c:\Users\letbounce\Desktop\Коксюк_Диплом.docx")
OUT = Path(r"c:\Users\letbounce\Desktop\Коксюк_Диплом.docx")
DIAG = Path(r"g:\AndroidStudioProjects\docs\diagrams")

FIG_W = Cm(16)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p


def add_para(doc, text, bold_first=False):
    p = doc.add_paragraph()
    if bold_first and "." in text:
        idx = text.index(".")
        run_b = p.add_run(text[: idx + 1])
        run_b.bold = True
        p.add_run(text[idx + 1 :])
    else:
        p.add_run(text)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_figure(doc, image_name, caption):
    img = DIAG / image_name
    if not img.exists():
        doc.add_paragraph(f"[ВСТАВИТИ РИСУНОК: {image_name}]")
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img), width=FIG_W)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.font.size = Pt(12)
    doc.add_paragraph()


def add_table_caption(doc, num, title):
    p = doc.add_paragraph(f"Таблиця {num}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.bold = True


def build_section3(doc):
    add_heading(doc, "РОЗДІЛ 3. ПРОЕКТУВАННЯ ТА РЕАЛІЗАЦІЯ КОМПОНЕНТІВ ПІДСИСТЕМИ", 1)

    intro = (
        "У третьому розділі викладено проектні рішення щодо інформаційного, технічного та "
        "програмного забезпечення підсистеми «Мобільний термінал водія» RoutePulse, а також "
        "результати реалізації прототипу на контрольному прикладі. Матеріал розділу побудовано "
        "відповідно до методичних рекомендацій КНЕУ (РД 50-34.698-90 для ІЗ і технічного "
        "забезпечення) та узгоджений з моделями, наведеними в розділі 2."
    )
    add_para(doc, intro)

    # ----- 3.1 -----
    add_heading(doc, "3.1. Інформаційне забезпечення", 2)

    add_para(
        doc,
        "Загальна характеристика інформаційного забезпечення. "
        "Інформаційне забезпечення (ІЗ) підсистеми RoutePulse охоплює сукупність методів і "
        "засобів формування, зберігання, контролю та видачі даних, необхідних для автоматизації "
        "робочих процесів водія міського пасажирського транспорту та диспетчерської служби АТП. "
        "Основні інформаційні об'єкти: облікові дані водіїв, довідники маршрутів і транспортних "
        "засобів, електронні дорожні листи (waybills), журнал інцидентів, бакети GPS-телеметрії, "
        "а також файли фотофіксації інцидентів. Носіями даних є центральна база MongoDB "
        "(колекції routepulse_db), файлове сховище /server/uploads для JPEG-фото, локальна "
        "черга Room на смартфоні водія та EncryptedSharedPreferences для JWT-токена.",
    )

    add_para(
        doc,
        "Принципи організації ІЗ: єдине джерело правди на сервері (MongoDB); REST API як "
        "канал обміну; offline-first для телеметрії; soft-delete та version_history для "
        "інцидентів; валідація на клієнті та сервері (express-validator, Mongoose enum). "
        "Загальна схема інформаційного забезпечення наведена на рисунку 3.1.",
    )
    add_figure(
        doc,
        "15-info-support-scheme.png",
        "Рисунок 3.1 — Загальна схема інформаційного забезпечення підсистеми «Мобільний термінал водія»",
    )

    add_para(
        doc,
        "На рисунку 3.1 відображено повний ланцюг: джерела інформації (водії, "
        "адміністратор/диспетчер, GPS, OpenStreetMap) → засоби збору (Compose-екрани, seed-скрипти, "
        "LocationTrackingService) → перетворення та контроль (валідація, JWT, класифікатори) → "
        "центральне сховище MongoDB (шість колекцій) → засоби видачі (REST API, PDF-експорт, "
        "live-карта) → споживачі (диспетчери, аналітик, служба безпеки руху).",
    )

    add_para(
        doc,
        "Організація збору і передачі первинної інформації. "
        "Первинна інформація надходить від водіїв через мобільний застосунок RoutePulse "
        "(~150 пристроїв), від адміністратора — через npm run seed, MongoDB Compass та smoke-тести, "
        "від зовнішніх служб — GPS-супутники та OSM tile-server. Періодичність: GPS-точки — кожні "
        "5 с під час активного рейсу; синхронізація батчів телеметрії — кожні 15 хв (WorkManager); "
        "інциденти та дорожні листи — за подією (on-demand). Формат обміну — JSON по HTTPS з "
        "заголовком Authorization: Bearer.",
    )

    add_para(
        doc,
        "Побудова системи класифікації та кодування. "
        "Для ідентифікації об'єктів використовуються: driver_id у форматі DRV-NNNN, vehicle_id — "
        "KP-NNNN, BSON ObjectId (24 hex) для документів MongoDB. Класифікатори статусів: "
        "waybill.status ∈ {assigned, in_progress, completed, cancelled}; incident.type ∈ "
        "{accident, breakdown, traffic_jam, other}; incident.status ∈ {open, resolved, completed}; "
        "deletion_reason_code — п'ять фіксованих значень (error, duplicate, wrong_data, "
        "equipment_issue, other). Кодування фото — JPEG у файловому сховищі, передача з клієнта — "
        "Base64 у полі photo_base64 JSON-запиту.",
    )

    add_para(
        doc,
        "Проектування форм первинних документів, машинограм та відеокадрів. "
        "У прототипі RoutePulse традиційні паперові форми замінено електронними екранами Jetpack "
        "Compose: LoginScreen (вхід), RouteDashboardScreen (старт рейсу), ActiveTripScreen "
        "(активний рейс), IncidentReportScreen (звіт про інцидент), IncidentsListScreen (журнал), "
        "TripMapHubScreen / TripMapRouteScreen (карта). Вихідні машинограми формуються як JSON-"
        "відповіді REST API та можуть експортуватися у PDF (PdfExportWriter).",
    )

    add_para(
        doc,
        "Структура інформаційних масивів. "
        "Основні логічні масиви відповідають колекціям MongoDB. Для колекції waybills "
        "характерні поля: driver_id (індекс), route_id (посилання), status, started_at, completed_at, "
        "deleted_at. Для incidents — waybill_id, driver_id, type, location (embedded), photo_url, "
        "version_history[]. Для telemetry — waybill_id, bucket_start, locations[] (до ~100 точок "
        "на батч), count. Детальна фізична структура наведена в даталогічній моделі (п. 3.1.6).",
    )

    add_para(
        doc,
        "Вибір СКБД. "
        "Обрано MongoDB 6.x (документо-орієнтована СКБД) з ODM Mongoose. Обґрунтування: "
        "гнучкі вкладені структури (масив зупинок у routes, бакети координат у telemetry, "
        "історія версій інцидентів); горизонтальне масштабування при зростанні телеметрії; "
        "швидка розробка прототипу на Node.js; відповідність JSON-формату REST API. "
        "Обмеження: відсутність жорстких JOIN — компенсується денормалізацією (route_number у waybill) "
        "та посиланнями ObjectId.",
    )

    add_para(
        doc,
        "Інфологічна модель бази даних. "
        "Інфологічна ER-модель побудована в нотації Crow's foot і відображає сутності предметної "
        "області без прив'язки до фізичної реалізації MongoDB. На рисунку 3.2 наведено зв'язки: "
        "DRIVER — «виконує» — WAYBILL; ROUTE — «обслуговує» / «містить» — STOP; VEHICLE — "
        "«експлуатується»; WAYBILL — «має» — INCIDENT; WAYBILL — «генерує треки» — TELEMETRY.",
    )
    add_figure(
        doc,
        "09-er-infological.png",
        "Рисунок 3.2 — Інфологічна модель бази даних інформаційної підсистеми (нотація Crow's foot)",
    )

    add_para(
        doc,
        "Даталогічна модель бази даних. "
        "Фізична схема реалізована у вигляді шести колекцій бази routepulse_db (рисунок 3.3). "
        "Колекції drivers та vehicles мають унікальні індекси за бізнес-ідентифікаторами. "
        "У routes зупинки зберігаються як embedded-масив stops[] без власного _id. "
        "Колекції waybills, incidents та telemetry містять посилання ObjectId на waybills/routes, "
        "індекси за driver_id та waybill_id для прискорення запитів диспетчерської аналітики.",
    )
    add_figure(
        doc,
        "10-er-datalogical.png",
        "Рисунок 3.3 — Даталогічна модель бази даних (фізична схема колекцій MongoDB)",
    )

    add_para(
        doc,
        "Контрольний приклад наповнення БД виконується скриптом server/scripts/seed.js: "
        "створюються тестові водії (зокрема DRV-1042 / password123), маршрути зі зупинками, "
        "транспортні засоби. Перевірка — через npm run smoke або MongoDB Compass після запуску "
        "рейсу та реєстрації інциденту з мобільного клієнта.",
    )

    # ----- 3.2 -----
    add_heading(doc, "3.2. Технічне забезпечення", 2)

    add_para(
        doc,
        "Загальні положення та схема автоматизації. "
        "Підсистема RoutePulse функціонує за архітектурою клієнт–сервер: мобільні термінали "
        "водіїв (~150 Android-пристроїв) взаємодіють із сервером застосунків (Node.js/Express на "
        "Linux VPS або private cloud), який зберігає дані в MongoDB. Зовнішні сервіси: "
        "супутникова GPS-навігація та OpenStreetMap tile-server для картографії. "
        "Загальна схема автоматизації наведена на рисунку 3.4.",
    )
    add_figure(
        doc,
        "16-automation-scheme.png",
        "Рисунок 3.4 — Загальна схема автоматизації (взаємодія Android · API-сервер · MongoDB)",
    )

    add_para(
        doc,
        "На рисунку 3.4 показано три рівні: mobile tier (Compose UI, ViewModels, Retrofit, Room, "
        "FGS, WorkManager), application server tier (Express, middleware requireAuth, controllers, "
        "Mongoose, /uploads), data tier (routepulse_db, WiredTiger). Канали: HTTPS REST+JWT між "
        "клієнтом і API; TCP 27017 між API і MongoDB (лише внутрішній LAN).",
    )

    add_para(
        doc,
        "Структура комплексу технічних засобів. "
        "Комплекс технічних засобів (КТЗ) включає серверну інфраструктуру, АРМ диспетчерської "
        "служби, мобільні термінали водіїв та мережеве обладнання (рисунок 3.5). "
        "Серверний вузол: Application Server (4 vCPU, 8 GB RAM, Ubuntu 22.04, Node.js 18, PM2) "
        "та Database Server (16 GB RAM, MongoDB 6.x, NVMe + RAID для бекапів). "
        "Резервне копіювання: NAS 2×2 TB, mongodump cron, архів /uploads.",
    )
    add_figure(
        doc,
        "17-ktz-structure.png",
        "Рисунок 3.5 — Структура комплексу технічних засобів (КТЗ) інформаційної підсистеми",
    )

    add_para(
        doc,
        "Диспетчерська служба обладнана 5–6 робочими станціями (Intel Core i5, 16 GB RAM, "
        "монітор 24″), 1–2 великими моніторинг-дисплеями 55″ 4K, мережевим принтером для звітів, "
        "UPS. Загальна кількість мобільних терміналів — ~165 (150 робочих + 10–15 резервних). "
        "Інженерне забезпечення: UPS серверної 3000 ВА, кондиціонування серверної 24±2 °C, "
        "резервний дизель-генератор АТП.",
    )

    add_para(
        doc,
        "Опис автоматизованого робочого місця водія. "
        "АРМ водія — переносний комплекс на базі смартфона Android 7.0+ (API 24), встановленого "
        "в кабіні транспортного засобу (рисунок 3.6). Мінімальні вимоги: 4 ядра CPU, 2 GB RAM, "
        "16 GB ROM, GPS L1, камера ≥5 МП, LTE/4G. Аксесуари: утримувач на лобовому склі, "
        "зарядка 12V→USB (2 А), кабель USB. Програмний стек деталізовано в п. 3.3.",
    )
    add_figure(
        doc,
        "18-driver-workstation.png",
        "Рисунок 3.6 — Схема автоматизованого робочого місця (АРМ) водія міського пасажирського транспорту",
    )

    add_para(
        doc,
        "Схема мережі передачі даних. "
        "Мережева архітектура побудована за принципом сегментації на три зони (рисунок 3.7): "
        "(1) зовнішня мережа — смартфони водіїв, оператори 4G/LTE, Internet, OSM; "
        "(2) периметр АТП (DMZ) — FTTH 100 Mbps, firewall (pfSense/OPNsense, TLS, NAT 443→API), "
        "WiFi 6 AP у депо (SSID RoutePulse-AT, WPA2-Enterprise); "
        "(3) внутрішня LAN 192.168.10.0/24 — L2 switch, app server :3000, db server :27017 "
        "(internal-only), АРМ диспетчерів, NAS, принтер.",
    )
    add_figure(
        doc,
        "19-network-scheme.png",
        "Рисунок 3.7 — Схема мережі передачі даних інформаційної підсистеми",
    )

    add_para(
        doc,
        "Об'ємно-часові характеристики (за розрахунками на рис. 3.7): при ~150 водіях, "
        "GPS 1 точка/5 с, батч 100 точок/15 хв — орієнтовно ~28 MB/добу телеметрії на fleet; "
        "транзакції: ~150 login, ~200 waybill, ~30–50 інцидентів/добу; фото інцидентів — "
        "~30–50 MB/добу. Канал ISP 100 Mbps забезпечує запас пропускної здатності для пікових "
        "навантажень на початку змін (07:00–09:00). MongoDB порт 27017 не експонується в Internet.",
    )

    # ----- 3.3 -----
    add_heading(doc, "3.3. Програмне забезпечення", 2)

    add_para(
        doc,
        "Структура програмного забезпечення. "
        "Програмне забезпечення RoutePulse класифіковано на системне, RAD/CASE, прикладне "
        "(Android + Backend) та програмну документацію (рисунок 3.8).",
    )
    add_figure(
        doc,
        "20-software-structure.png",
        "Рисунок 3.8 — Структура програмного забезпечення інформаційної підсистеми (RoutePulse)",
    )

    add_para(
        doc,
        "Системне програмне забезпечення. "
        "Сервер: Ubuntu Server 22.04 LTS, Node.js 18 LTS, MongoDB 6.x (WiredTiger), PM2/systemd, "
        "nginx (reverse-proxy, TLS, gzip), UFW, Fail2ban, mongodump+cron. "
        "Клієнт: Android 7.0+ (target SDK 36), ART, Google Play Services (Location, Camera), "
        "Material Components 3. Сервісне ПЗ: моніторинг (Prometheus/Grafana або Uptime Kuma), "
        "резервне копіювання на NAS.",
    )

    add_para(
        doc,
        "RAD-засоби та CASE. "
        "Розробка: Android Studio (Kotlin, Compose), Cursor/VS Code (Node.js), Gradle 8 + KSP, "
        "npm. Моделювання: SVG-діаграми (Inkscape/draw.io), PlantUML; MongoDB Compass для "
        "перегляду БД. Тестування: JUnit, Espresso, Compose UI Test, smoke.js, smoke-curl.sh. "
        "Керування версіями: Git, GitHub/GitLab.",
    )

    add_para(
        doc,
        "Прикладне програмне забезпечення. "
        "Android-клієнт RoutePulse реалізовано за Clean Architecture + MVVM: шари presentation "
        "(Compose-екрани, ViewModel), domain (моделі, UseCase, інтерфейси репозиторіїв), data "
        "(Retrofit ApiService, RepositoryImpl, Room LocalQueue, DataStore, EncryptedSharedPreferences). "
        "Фонові компоненти: LocationTrackingService (FGS, інтервал ~5 с), TelemetrySyncWorker "
        "(PeriodicWorkRequest, 15 хв, max 100 точок). "
        "Серверна частина: Express 4, маршрути /api/auth, /api/routes, /api/waybills, /api/incidents, "
        "/api/telemetry, /api/map, /api/vehicles; контролери auth, waybill, incident, telemetry, "
        "route, map, vehicle; bcrypt (cost 10), jsonwebtoken (HS256, 8 год), express-validator.",
    )

    add_para(
        doc,
        "Програмна документація. "
        "До складу документації входять: README.md (корінь і server/), docs/*.md (UX, demo, smoke), "
        "docs/diagrams/ (20 SVG+PNG), KDoc-коментарі Kotlin, JSDoc-стиль у Node.js. "
        "Артефакти збірки: routepulse-release.apk / .aab. Мінімальна конфігурація АРМ водія "
        "відповідає рисунку 3.6; розгортання сервера — npm install, .env, npm run seed, npm start.",
    )

    # ----- 3.4 -----
    add_heading(doc, "3.4. Результати реалізації інформаційної підсистеми", 2)

    add_para(
        doc,
        "У рамках бакалаврського проекту реалізовано робочий прототип інформаційної підсистеми "
        "RoutePulse, що включає Android-застосунок (модуль app) та серверну підсистему (каталог "
        "server). Оригінальні проектні рішення полягають у поєднанні: (1) мобільного термінала "
        "водія на Jetpack Compose з offline-чергою телеметрії; (2) REST API на Node.js з JWT-"
        "автентифікацією; (3) документо-орієнтованої моделі MongoDB з підтримкою версійності "
        "інцидентів та soft-delete дорожніх листів.",
    )

    add_para(
        doc,
        "Контрольний приклад апробації. "
        "Сценарій демонстрації (5–7 хв): (1) запуск бекенду — cd server, npm run seed, npm start; "
        "(2) авторизація в Android-клієнті — DRV-1042 / password123; (3) перегляд маршрутів на "
        "RouteDashboardScreen, старт дорожнього листа; (4) активний рейс — фоновий GPS-трекінг; "
        "(5) реєстрація інциденту з фото на IncidentReportScreen; (6) завершення рейсу "
        "PATCH /api/waybills/{id}/complete; (7) перевірка записів у MongoDB (колекції waybills, "
        "incidents, telemetry). Автоматизована перевірка API — npm run smoke та bash scripts/smoke-curl.sh.",
    )

    add_para(
        doc,
        "Отримані результати підтверджують досягнення цілей, сформульованих у розділі 2: "
        "реалізовано FR-01…FR-10 та NFR-01…NFR-04 (див. діаграму трасування, рис. 2.10). "
        "Прототип придатний для пілотного впровадження на АТП масштабу ~150 водіїв за умови "
        "розгортання серверного вузла, налаштування HTTPS та політик безпеки мобільних пристроїв. "
        "Подальший розвиток: окремий веб-кабінет диспетчера, push-сповіщення (FCM), "
        "мультипарт-завантаження фото, горизонтальне масштабування MongoDB.",
    )


def main():
    if not THESIS.exists():
        raise SystemExit(f"File not found: {THESIS}")
    doc = Document(str(THESIS))
    build_section3(doc)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
