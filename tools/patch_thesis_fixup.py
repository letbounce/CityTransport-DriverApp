# -*- coding: utf-8 -*-
"""Доробка після patch_thesis_kbr2026.py."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

THESIS = Path(r"d:\Downloads\Коксюк_Диплом.docx")
EXTRA_BIB = Path(__file__).resolve().parent / "patch_thesis_kbr2026.py"


def delete_paragraph(paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def is_toc(p) -> bool:
    return p.style and p.style.name.startswith("toc")


def remove_section(doc: Document, heading_prefix: str, until_re: str) -> None:
    to_del = []
    active = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if is_toc(p):
            if heading_prefix in t and t.startswith(heading_prefix.split()[0][:3]):
                to_del.append(p)
            continue
        if t.startswith(heading_prefix):
            active = True
            to_del.append(p)
            continue
        if active:
            if re.match(until_re, t):
                break
            to_del.append(p)
    for p in to_del:
        delete_paragraph(p)


def fix_reference_316(doc: Document) -> None:
    for p in doc.paragraphs:
        if "п. 3.1.6" in p.text or "п.3.1.6" in p.text:
            p.text = p.text.replace("п. 3.1.6", "п. 3.1.5").replace("п.3.1.6", "п.3.1.5")


def rebuild_bib_only(doc: Document) -> None:
    import importlib.util

    spec = importlib.util.spec_from_file_location(
        "patch_kbr", Path(__file__).parent / "patch_thesis_kbr2026.py"
    )
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    extra_bib = mod.EXTRA_BIB

    entries = []
    after = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "ПЕРЕЛІК ВИКОРИСТАНИХ ДЖЕРЕЛ":
            after = True
            continue
        if after:
            if t.startswith("ДОДАТКИ"):
                break
            if t:
                t = re.sub(r"^\d+\.\s*", "", t)
                entries.append(t)
    seen = set()
    merged = []
    for e in entries:
        k = e[:55]
        if k not in seen:
            seen.add(k)
            merged.append(e)
    for e in extra_bib:
        if len(merged) >= 32:
            break
        k = e[:55]
        if k not in seen:
            seen.add(k)
            merged.append(e)
    merged = merged[:32]

    heading = None
    dod = None
    for p in doc.paragraphs:
        if p.text.strip() == "ПЕРЕЛІК ВИКОРИСТАНИХ ДЖЕРЕЛ":
            heading = p
        if p.text.strip() == "ДОДАТКИ":
            dod = p
    if not heading or not dod:
        return

    to_del = []
    grab = False
    for p in doc.paragraphs:
        if p is heading:
            grab = True
            continue
        if grab:
            if p is dod:
                break
            if p.text.strip():
                to_del.append(p)
    for p in to_del:
        delete_paragraph(p)

    for i, e in enumerate(merged, 1):
        np = dod.insert_paragraph_before(f"{i}. {e}")
        np.paragraph_format.line_spacing = 1.5
        try:
            np.style = "Normal"
        except KeyError:
            pass

    for p in doc.paragraphs:
        if "список літератури з" in p.text:
            p.text = re.sub(r"\d+ найменувань", "32 найменувань", p.text)


def add_more_citations(doc: Document) -> None:
    """Додати [N] за номером у списку джерел."""
    bib_map: dict[str, int] = {}
    after = False
    n = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "ПЕРЕЛІК ВИКОРИСТАНИХ ДЖЕРЕЛ":
            after = True
            continue
        if after:
            if t.startswith("ДОДАТКИ"):
                break
            m = re.match(r"^(\d+)\.\s+", t)
            if m:
                n = int(m.group(1))
                low = t.lower()
                for key in ("wialon", "fleet complete", "mongodb", "iso/iec/ieee 29148", "iso/iec 19501",
                            "дсту 19.701", "рд 50-34", "kotlin", "hoang", "truica", "ситник", "express",
                            "workmanager", "room", "retrofit", "osmdroid", "avtogps", "jwt", "owasp"):
                    if key in low and key not in bib_map:
                        bib_map[key] = n

    replacements = [
        (r"\bWialon\b(?!\s*\[\d+\])", bib_map.get("wialon", 11)),
        (r"\bFleet Complete\b(?!\s*\[\d+\])", bib_map.get("fleet complete", 12)),
        (r"\bMongoDB\b(?!\s*\[\d+\])", bib_map.get("mongodb", 4)),
        (r"ISO/IEC/IEEE 29148:2018(?!\s*\[\d+\])", bib_map.get("iso/iec/ieee 29148", 8)),
        (r"ISO/IEC 19501:2005(?!\s*\[\d+\])", bib_map.get("iso/iec 19501", 7)),
        (r"ДСТУ 19\.701-90(?!\s*\[\d+\])", bib_map.get("дсту 19.701", 10)),
        (r"РД 50-34\.698-90(?!\s*\[\d+\])", bib_map.get("рд 50-34", 10)),
        (r"\bKotlin\b(?!\s*\[\d+\])", bib_map.get("kotlin", 30)),
        (r"Jetpack Compose(?!\s*\[\d+\])", bib_map.get("jetpack", 3)),
        (r"WorkManager(?!\s*\[\d+\])", bib_map.get("workmanager", 3)),
        (r"\bRoom\b(?!\s*\[\d+\])", bib_map.get("room", 31)),
        (r"AvtoGPS(?!\s*\[\d+\])", bib_map.get("avtogps", 27)),
    ]
    in_bib = False
    for p in doc.paragraphs:
        t = p.text
        if t.strip() == "ПЕРЕЛІК ВИКОРИСТАНИХ ДЖЕРЕЛ":
            in_bib = True
            continue
        if in_bib:
            if t.strip().startswith("ДОДАТКИ"):
                in_bib = False
            continue
        if is_toc(p):
            continue
        for pat, num in replacements:
            if num and re.search(pat, t):
                t = re.sub(pat, lambda m, n=num: f"{m.group(0)} [{n}]", t, count=1)
        p.text = t


def main() -> None:
    doc = Document(str(THESIS))
    remove_section(doc, "3.1.6.", r"^3\.2\.")
    remove_section(doc, "3.5.", r"^ВИСНОВКИ")
    fix_reference_316(doc)
    rebuild_bib_only(doc)
    add_more_citations(doc)
    doc.save(str(THESIS))
    print("fixup done")


if __name__ == "__main__":
    main()
