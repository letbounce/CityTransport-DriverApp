# -*- coding: utf-8 -*-
"""Append Додатки and Графічні матеріали with screenshots to thesis docx."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.enum.section import WD_SECTION

THESIS = Path(r"c:\Users\letbounce\Desktop\Коксюк_Диплом.docx")
SRC = Path(r"c:\Users\letbounce\Desktop\Дипломка")
FIG_W = Cm(15)

# (filename in Дипломка, caption, optional explanation)
APPENDIX_ITEMS = [
    (
        "Waybill.png",
        "Рисунок А.1 — Фрагмент серверної моделі даних: схема Mongoose для колекції waybills (файл Waybill.js)",
        "На рисунку А.1 наведено визначення полів дорожнього листа: посилання на водія (driver_id), "
        "маршрут (route_id), статус (enum: assigned, in_progress, completed, cancelled), часові мітки "
        "та поля soft-delete (deleted_at, deletion_reason_code). Схема узгоджена з даталогічною моделлю "
        "на рисунку 3.3.",
    ),
    (
        "Incident.png",
        "Рисунок А.2 — Фрагмент серверної моделі даних: схема Mongoose для колекції incidents (файл Incident.js)",
        "На рисунку А.2 показано структуру документа інциденту: тип події (accident, breakdown, traffic_jam, other), "
        "вбудований об'єкт location, photo_url, статус, масив version_history для зберігання історії змін "
        "та атрибут is_modified. Реалізація відповідає вимогам FR-06 та FR-07.",
    ),
]

GRAPHIC_ITEMS = [
    (
        "Вхід.png",
        "Рисунок Г.1 — Екран авторизації водія (LoginScreen)",
        "Початковий екран підсистеми «Мобільний термінал водія»: введення ідентифікатора DRV-NNNN "
        "та пароля, перевірка на сервері (bcrypt, JWT). Відповідає FR-01.",
    ),
    (
        "Головне меню.png",
        "Рисунок Г.2 — Головне меню після успішної авторизації (HomeMenuScreen)",
        "Навігаційний хаб застосунку: перехід до дорожніх листів, інцидентів, карти маршруту та виходу.",
    ),
    (
        "Створення дорожнього листа.png",
        "Рисунок Г.3 — Екран створення дорожнього листа (RouteDashboardScreen)",
        "Вибір маршруту №114, транспортного засобу та перегляд розкладу зупинок перед початком рейсу. FR-02, FR-03.",
    ),
    (
        "Лист рейсів.png",
        "Рисунок Г.4 — Перелік дорожніх листів / активних рейсів водія",
        "Відображення створених waybill із можливістю переходу до активного рейсу та завершення.",
    ),
    (
        "Редагування дорожнього листа.png",
        "Рисунок Г.5 — Редагування параметрів дорожнього листа",
        "Зміна даних waybill під час виконання рейсу (узгоджено з API PATCH /api/waybills/:id).",
    ),
    (
        "Мапа маршруту.png",
        "Рисунок Г.6 — Карта маршруту №114 із зупинками (TripMapRouteScreen, OSMdroid)",
        "Візуалізація траси та зупинок на тайлах OpenStreetMap. Відповідає FR-08.",
    ),
    (
        "Створення звіту про інцидент.png",
        "Рисунок Г.7 — Реєстрація звіту про інцидент (IncidentReportScreen)",
        "Фіксація часу, GPS-координат, ознаки руху та опційного фото. FR-06.",
    ),
    (
        "Редагування інциденту.png",
        "Рисунок Г.8 — Редагування зареєстрованого інциденту",
        "Коригування полів інциденту з фіксацією в version_history. FR-07.",
    ),
    (
        "Архівування інциденту.png",
        "Рисунок Г.9 — Архівування інциденту (soft-delete)",
        "Вказання причини архівації та збереження запису з deleted_at у MongoDB.",
    ),
    (
        "Архів інцидентів.png",
        "Рисунок Г.10 — Перегляд архіву інцидентів (IncidentsListScreen)",
        "Список заархівованих звітів із можливістю відновлення або перегляду.",
    ),
    (
        "waybill_completed.png",
        "Рисунок Г.11 — Документ колекції waybills у MongoDB після завершення рейсу",
        "Контрольний приклад: waybill маршруту №114, статус completed, DRV-1042 (п. 3.4).",
    ),
]

APPENDIX_B_ITEMS = [
    (
        "Приклад звіту інциденту в PDF.png",
        "Рисунок А.3 — Приклад сформованого PDF-звіту про інцидент",
        "Експорт звіту для служби безпеки руху та диспетчерської служби (форматування згідно з вимогами АТП).",
    ),
]


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p


def add_figure_block(doc, image_path: Path, caption: str, explanation: str | None = None):
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=FIG_W)
    else:
        doc.add_paragraph(f"[Відсутній файл: {image_path.name}]")

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.font.size = Pt(12)

    if explanation:
        add_body(doc, explanation)
    doc.add_paragraph()


def add_page_break(doc):
    doc.add_page_break()


def build_appendices(doc):
    add_page_break(doc)
    doc.add_heading("ДОДАТКИ", level=1)

    add_body(
        doc,
        "У додатках наведено фрагменти програмної реалізації серверної частини підсистеми RoutePulse, "
        "що доповнюють опис прикладного програмного забезпечення (розділ 3.3) та даталогічної моделі "
        "бази даних (розділ 3.1).",
    )

    doc.add_heading("Додаток А", level=2)
    doc.add_heading(
        "Фрагменти моделей даних серверної частини (Node.js / Mongoose)",
        level=3,
    )

    for fname, caption, explanation in APPENDIX_ITEMS:
        add_figure_block(doc, SRC / fname, caption, explanation)

    doc.add_heading("Додаток Б", level=2)
    doc.add_heading("Приклад вихідного документа звіту про інцидент", level=3)
    for fname, caption, explanation in APPENDIX_B_ITEMS:
        add_figure_block(doc, SRC / fname, caption, explanation)


def build_graphic_materials(doc):
    add_page_break(doc)
    doc.add_heading("ГРАФІЧНІ МАТЕРІАЛИ", level=1)

    add_body(
        doc,
        "У графічних матеріалах наведено екранні форми мобільного застосунку RoutePulse та результат "
        "збереження даних у MongoDB за результатами апробації на контрольному прикладі (DRV-1042, "
        "маршрут №114). Матеріали доповнюють UML-діаграми розділу 2 та схеми розділу 3 (рис. Г.1–Г.11) "
        "і ілюструють повний сценарій роботи водія: від авторизації до архівування інцидентів.",
    )

    doc.add_heading(
        "Інтерфейс мобільного терміналу водія та результати апробації",
        level=2,
    )

    for fname, caption, explanation in GRAPHIC_ITEMS:
        add_figure_block(doc, SRC / fname, caption, explanation)


def already_appended(doc) -> bool:
    return any(
        p.text.strip() in ("ДОДАТКИ", "ГРАФІЧНІ МАТЕРІАЛИ")
        for p in doc.paragraphs
        if p.style.name.startswith("Heading")
    )


def main():
    if not THESIS.exists():
        raise SystemExit(f"Не знайдено: {THESIS}")
    doc = Document(str(THESIS))
    if already_appended(doc):
        raise SystemExit(
            "Розділи «Додатки» / «Графічні матеріали» вже є в документі. "
            "Видаліть їх у Word і запустіть скрипт знову."
        )
    build_appendices(doc)
    build_graphic_materials(doc)
    doc.save(str(THESIS))
    print(f"Оновлено: {THESIS}")


if __name__ == "__main__":
    main()
