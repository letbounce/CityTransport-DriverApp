# -*- coding: utf-8 -*-
"""Перенесення з курсової в диплом матеріалу, адаптованого під RoutePulse."""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

THESIS = Path(r"c:\Users\letbounce\Desktop\Коксюк_Диплом.docx")
BACKUP = Path(r"c:\Users\letbounce\Desktop\Коксюк_Диплом_backup_course_merge.docx")


def style_table(table) -> None:
    try:
        table.style = "Table Grid"
    except KeyError:
        pass


def insert_body_before(anchor, text: str) -> None:
    p = anchor.insert_paragraph_before("")
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)


def insert_center_before(anchor, text: str) -> None:
    p = anchor.insert_paragraph_before(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def insert_table_before(anchor, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    insert_center_before(anchor, caption)
    doc = anchor.part.document
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    style_table(tbl)
    for j, h in enumerate(headers):
        tbl.rows[0].cells[j].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            tbl.rows[ri + 1].cells[ci].text = val
    anchor._element.addprevious(tbl._tbl)
    anchor.insert_paragraph_before("")


def find_paragraph(doc: Document, needle: str, start: int = 0):
    for p in doc.paragraphs[start:]:
        if needle in p.text:
            return p
    return None


def already_has(doc: Document, marker: str) -> bool:
    return any(marker in p.text for p in doc.paragraphs)


def fix_figure_numbering(doc: Document) -> None:
    """Усунути дубль: інфомодель 2.2, алгоритм і наступні +1."""
    pairs = [
        ("Рисунок 2.10 — Діаграма трасування", "Рисунок 2.11 — Діаграма трасування"),
        ("Діаграма трасування (рисунок 2.10)", "Діаграма трасування (рисунок 2.11)"),
        ("рис. 2.10", "рис. 2.11"),
        ("Рисунок 2.9 — Діаграма класів", "Рисунок 2.10 — Діаграма класів"),
        ("діаграмою класів UML (рисунок 2.9)", "діаграмою класів UML (рисунок 2.10)"),
        ("Діаграма класів (рисунок 2.9)", "Діаграма класів (рисунок 2.10)"),
        ("Рисунок 2.8 — Діаграма діяльності", "Рисунок 2.9 — Діаграма діяльності"),
        ("Рисунок 2.8 відображає", "Рисунок 2.9 відображає"),
        ("Рисунок 2.7 — Діаграма послідовності: передача GPS", "Рисунок 2.8 — Діаграма послідовності: передача GPS"),
        ("Рисунок 2.7 відображає", "Рисунок 2.8 відображає"),
        ("Рисунок 2.6 — Діаграма послідовності: реєстрація інциденту", "Рисунок 2.7 — Діаграма послідовності: реєстрація інциденту"),
        ("Рисунок 2.6 ілюструє", "Рисунок 2.7 ілюструє"),
        ("Рисунок 2.5 — Діаграма послідовності: створення", "Рисунок 2.6 — Діаграма послідовності: створення"),
        ("Рисунок 2.5 описує", "Рисунок 2.6 описує"),
        ("Рисунок 2.4 — Діаграма послідовності: авторизація", "Рисунок 2.5 — Діаграма послідовності: авторизація"),
        ("Рисунок 2.4 відображає сценарій авторизації", "Рисунок 2.5 відображає сценарій авторизації"),
        ("Рисунок 2.3 — Діаграма прецедентів", "Рисунок 2.4 — Діаграма прецедентів"),
        ("Діаграма прецедентів (рисунок 2.3)", "Діаграма прецедентів (рисунок 2.4)"),
        ("Рисунок 2.2 — Схема алгоритму", "Рисунок 2.3 — Схема алгоритму"),
        ("Алгоритм (рисунок 2.2) розпочинається", "Алгоритм (рисунок 2.3) розпочинається"),
        ("на рисунку 2.2. Алгоритм", "на рисунку 2.3. Алгоритм"),
    ]
    for p in doc.paragraphs:
        for old, new in pairs:
            if old in p.text:
                p.text = p.text.replace(old, new)


def merge_section11_as_is(doc: Document) -> None:
    marker = "Single Source of Truth"
    if already_has(doc, marker):
        return
    anchor = find_paragraph(doc, "Ключовою проблемою існуючих систем управління")
    if not anchor:
        return
    insert_body_before(
        anchor,
        "Аналіз поточного стану (AS-IS) типового АТП показує фрагментарну автоматизацію: "
        "GPS-моніторинг, облік у MS Excel або 1С, паперові дорожні листи та голосовий зв'язок "
        "водія з диспетчером не об'єднані в єдиному інформаційному просторі. Виникає розрив між "
        "оперативним (диспетчеризація), технічним (СТО) та фінансовим (бухгалтерія) контурами, "
        "що ускладнює оперативні рішення та підвищує частку ручного введення даних. "
        "Підсистема RoutePulse усуває цей розрив для контурів водія та диспетчерської служби, "
        "формуючи єдине джерело достовірних даних про рейс, телеметрію та інциденти "
        f"({marker}).",
    )


def merge_section12_analogs(doc: Document) -> None:
    if already_has(doc, "UMT (Система управління транспортом)"):
        return
    anchor = find_paragraph(doc, "Серед практичних систем на ринку найбільш поширеними")
    if not anchor:
        return
    insert_body_before(
        anchor,
        "Для об'єктивного порівняння відібрано системи, релевантні задачі мобільної "
        "диспетчеризації: Wialon (Gurtam) — еталон GPS-моніторингу з модулями NimBus "
        "(пасажирські перевезення) та Fleetrun (ТО); UMT — український комплекс з модулем "
        "пасажирських перевезень та інтеграцією з 1С; Benish GPS — моніторинг і контроль палива "
        "на вітчизняному ринку; Fleet Complete та AvtoGPS — хмарні/локальні рішення телематики. "
        "Критерії оцінювання сформовано за «вузькими місцями» AS-IS: GPS-трекінг, мобільний "
        "термінал водія, електронний дорожній лист, реєстрація інцидентів, офлайн-телеметрія, "
        "відкритий API та вартість впровадження для АТП ~150 ТЗ.",
    )

    gap_anchor = find_paragraph(doc, "Аналіз таблиці 1.1 свідчить")
    if gap_anchor and not already_has(doc, "прогалину ринку"):
        insert_body_before(
            gap_anchor,
            "Порівняльний аналіз (за аналогією з курсовим проєктом повної ІС управління "
            "рухомим складом) виявляє прогалину ринку: комерційні платформи часто вимагають "
            "окремих модулів (NimBus, Fleetrun) або спеціалізованого бортового обладнання, "
            "тоді як RoutePulse концентрує ключовий функціонал мобільного терміналу водія "
            "(waybill, GPS, інцидент з фото, offline-черга) у self-hosted прототипі на базі "
            "смартфона Android без ліцензійних платежів SaaS.",
        )


def merge_section211_stakeholders(doc: Document) -> None:
    if already_has(doc, "Стейкхолдери підсистеми"):
        return
    anchor = find_paragraph(doc, "БЦ-04:")
    if not anchor:
        anchor = find_paragraph(doc, "2.1.2. Специфікація функціональних")
    if not anchor:
        return
    # insert after BC-04 paragraph
    for p in doc.paragraphs:
        if "БЦ-04:" in p.text:
            anchor = p
            break
    insert_body_before(
        anchor,
        "Стейкхолдери підсистеми RoutePulse та їхні інтереси (адаптовано з аналізу "
        "зацікавлених сторін повної ІС АТП): директор — зниження операційних витрат і KPI; "
        "старший диспетчер зміни — оперативний контроль рейсів і інцидентів; диспетчер з "
        "безпеки руху — журнал інцидентів із фотодоказами; водій (~150 осіб) — спрощення "
        "електронного дорожнього листа та фіксація подій у рейсі; адміністратор ІС — "
        "розгортання API/MongoDB та оновлення мобільного клієнта.",
    )


def merge_section213_nfr_details(doc: Document) -> None:
    if already_has(doc, "Характеристика груп нефункціональних вимог"):
        return
    anchor = find_paragraph(doc, "Рисунок 2.1 — Діаграма нефункціональних")
    if not anchor:
        return
    texts = [
        "Характеристика груп нефункціональних вимог (узгоджено з аналізом умов експлуатації "
        "мобільного терміналу):",
        "Продуктивність — GPS-точка кожні 5 с під час активного waybill; відображення на "
        "карті диспетчера (майбутній веб-клієнт) — оновлення не рідше ніж за 15 хв або "
        "після накопичення 100 точок у батчі.",
        "Безпека — JWT (HS256, 8 год), bcrypt для паролів, EncryptedSharedPreferences на "
        "клієнті; у продакшені — лише HTTPS.",
        "Надійність та офлайн — Room-черга телеметрії, WorkManager з повторними спробами; "
        "відмова мережі не призводить до втрати GPS-даних активного рейсу.",
        "Ергономіка (Usability) — інтерфейс Compose для роботи в кабіні ТЗ: шрифт ≥16 sp, "
        "кнопки ≥64 dp, мінімум кроків для реєстрації інциденту (NFR-03).",
        "Масштабованість — архітектура розрахована на ~150 одночасних мобільних клієнтів "
        "та горизонтальне масштабування Node.js/MongoDB.",
        "Супровідність — відкритий REST API, smoke-тести (14 ендпоінтів), документація в "
        "репозиторії (README, протоколи інтеграційного тестування).",
        "Сумісність — клієнт Android 7.0+; сервер Ubuntu/Node.js 18+; обмін JSON; "
        "інтеграція з зовнішніми ІС «Кадри» та «Диспетчерська» через REST.",
    ]
    for t in reversed(texts):
        insert_body_before(anchor, t)


def merge_section221_task_details(doc: Document) -> None:
    if already_has(doc, "Періодичність виконання задачі"):
        return
    anchor = find_paragraph(doc, "Умови припинення автоматизованого розв'язання")
    if not anchor:
        return
    texts = [
        "Періодичність виконання задачі: оперативний контур — безперервний GPS-трекінг "
        "під час рейсу (in_progress); тактичний — створення/завершення waybill на зміну; "
        "аналітичний — агрегація телеметрії та архів інцидентів за запитом диспетчера/аналітика.",
        "Міжсистемні зв'язки: на вході — довідники водіїв, маршрутів і ТЗ (ІС «Кадри», "
        "«Диспетчерська»); на виході — дані для бухгалтерії (виконані рейси) та KPI "
        "(телеметрія), що відповідає інформаційній моделі (рис. 2.2).",
    ]
    for t in reversed(texts):
        insert_body_before(anchor, t)

    if not already_has(doc, "структурних одиниць ключових вихідних"):
        wb_anchor = find_paragraph(doc, "Таблиця 2.2")
        if wb_anchor:
            insert_body_before(
                wb_anchor,
                "Опис структурних одиниць ключових вихідних повідомлень (за зразком "
                "курсового проєкту, адаптовано під RoutePulse). Дорожній лист (WAYBILL_COMPL): "
                "ідентифікатор waybill, driver_id, route_id/route_number, vehicle_id, "
                "status=completed, started_at, completed_at. Звіт про інцидент (INCIDENT_RPT): "
                "type, description, stop_label, location{lat,lng}, photo_url, reported_at, "
                "waybill_id.",
            )


def merge_section231_uc_scenario(doc: Document) -> None:
    if already_has(doc, "Таблиця 2.7"):
        return
    anchor = find_paragraph(doc, "UC-06 «Завершити/архівувати рейс»")
    if not anchor:
        anchor = find_paragraph(doc, "Рисунок 2.4 — Діаграма прецедентів")
    if not anchor:
        return
    insert_body_before(
        anchor,
        "Для формалізації логіки ключових прецедентів розроблено текстовий сценарій "
        "(за зразком курсового проєкту) — див. табл. 2.7.",
    )
    insert_table_before(
        anchor,
        "Таблиця 2.7\nСценарій прецеденту «Реєстрація інциденту з фотофіксацією»",
        ["Елемент / Крок", "Опис", "Примітка"],
        [
            ["ID / Назва", "UC-05: Реєстрація інциденту з фото", ""],
            ["Актор", "Водій", ""],
            ["Передумови", "Активний waybill, авторизація JWT", ""],
            [
                "Основний потік",
                "1–5: форма інциденту → валідація → API → MongoDB → UI",
                "Див. рис. 2.7",
            ],
            [
                "Альтернативи",
                "Валідація на клієнті; помилки 4xx/5xx",
                "FR-06, NFR-02",
            ],
        ],
    )


def merge_section342_test_cases(doc: Document) -> None:
    if already_has(doc, "Таблиця 3.7"):
        return
    anchor = find_paragraph(doc, "Усі 14 API-ендпоінтів пройшли smoke-тестування")
    if not anchor:
        anchor = find_paragraph(doc, "3.4.2. Результати апробації")
    if not anchor:
        return
    insert_body_before(
        anchor,
        "Для верифікації функціональних вимог (за аналогією з тест-кейсами курсового "
        "проєкту) виконано сценарії взаємодії з прототипом RoutePulse — табл. 3.7.",
    )
    rows = [
        ["TC-01", "UC-01 / FR-01", "Login DRV-1042", "JWT, перехід на Home", "Пройдено"],
        ["TC-02", "UC-02 / FR-03", "Створення waybill м.114", "HTTP 201, in_progress", "Пройдено"],
        ["TC-03", "UC-04 / FR-09", "GPS під час рейсу", "Запис у Room, sync batch", "Пройдено"],
        ["TC-04", "UC-05 / FR-06", "Інцидент + фото", "HTTP 201, photo_url у БД", "Пройдено"],
        ["TC-05", "UC-06 / FR-04", "Завершення рейсу", "status=completed", "Пройдено"],
        ["TC-06", "UC-03 / FR-08", "Карта маршруту OSM", "GeoJSON, зупинки", "Пройдено"],
    ]
    insert_table_before(
        anchor,
        "Таблиця 3.7\nТест-кейси апробації підсистеми RoutePulse (контрольний приклад DRV-1042)",
        ["ID", "Вимога / UC", "Дія", "Очікуваний результат", "Статус"],
        rows,
    )


def extend_table_11(doc: Document) -> None:
    """Додаткова таблиця порівняння з курсовими аналогами."""
    if already_has(doc, "Таблиця 1.2"):
        return
    if already_has(doc, "табл. 1.2"):
        return
    anchor = find_paragraph(doc, "Таблиця 1.1")
    if not anchor:
        return
    insert_body_before(
        anchor,
        "Додатково до табл. 1.1 наведено порівняння з системами, проаналізованими "
        "у курсовому проєкті (повна ІС управління рухомим складом) — табл. 1.2.",
    )
    rows = [
        ["GPS-моніторинг", "Відмінно", "Відмінно", "Добре", "Так (5 с)"],
        ["Модуль пасажирського транспорту", "NimBus", "Спец. модуль", "Обмежено", "Маршрути+зупинки в RoutePulse"],
        ["Мобільний додаток водія", "WiaTag", "Так", "Так", "Android RoutePulse"],
        ["Відкритий API", "SDK/API", "REST, 1С", "API", "Повний REST"],
        ["Вартість для АТП 150 ТЗ", "Висока SaaS", "Середня/висока", "Середня", "Низька (self-hosted)"],
    ]
    insert_table_before(
        anchor,
        "Таблиця 1.2\nПорівняння з аналогами (фрагмент з курсового проєкту, адаптовано)",
        ["Критерій", "Wialon", "UMT", "Benish GPS", "RoutePulse"],
        rows,
    )


def main() -> None:
    if not THESIS.exists():
        raise SystemExit(f"Не знайдено: {THESIS}")
    shutil.copy2(THESIS, BACKUP)
    print(f"Резервна копія: {BACKUP}")

    doc = Document(str(THESIS))
    fix_figure_numbering(doc)
    merge_section11_as_is(doc)
    merge_section12_analogs(doc)
    extend_table_11(doc)
    merge_section211_stakeholders(doc)
    merge_section213_nfr_details(doc)
    merge_section221_task_details(doc)
    merge_section231_uc_scenario(doc)
    merge_section342_test_cases(doc)
    doc.save(str(THESIS))
    print(f"Оновлено: {THESIS}")


if __name__ == "__main__":
    main()
