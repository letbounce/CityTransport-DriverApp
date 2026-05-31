# -*- coding: utf-8 -*-
"""Перенести текст під заголовки та заповнити порожні підрозділи."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

THESIS = Path(r"c:\Users\letbounce\Desktop\Коксюк_Диплом.docx")
OUT = Path(r"c:\Users\letbounce\Desktop\Коксюк_Диплом_filled.docx")


def delete_paragraph(paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def set_heading(paragraph, level: int) -> None:
    paragraph.style = f"Heading {level}"


def add_body_after(anchor, text: str):
    nxt_el = anchor._element.getnext()
    if nxt_el is not None:
        from docx.text.paragraph import Paragraph

        p = Paragraph(nxt_el, anchor._parent).insert_paragraph_before(text)
    else:
        p = anchor._parent.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_bullet_after(anchor, text: str):
    nxt_el = anchor._element.getnext()
    if nxt_el is not None:
        from docx.text.paragraph import Paragraph

        p = Paragraph(nxt_el, anchor._parent).insert_paragraph_before("– " + text)
    else:
        p = anchor._parent.add_paragraph("– " + text)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    return p


def move_elements_after(heading, body_paras: list) -> None:
    for para in reversed(body_paras):
        heading._element.addnext(para._element)


def find_heading(doc, title_prefix: str):
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith(title_prefix) and not p.style.name.startswith("toc"):
            if len(t) < 120:
                return p
    return None


def collect_before(doc, heading, markers: tuple[str, ...], max_n: int = 12):
    paras = []
    idx = None
    hp = heading.text.strip()[:30]
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(hp[:10]) and p._element is heading._element:
            idx = i
            break
    if idx is None:
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == heading.text.strip():
                idx = i
                break
    if idx is None:
        return []
    for j in range(idx - 1, max(idx - max_n, -1), -1):
        t = doc.paragraphs[j].text.strip()
        if not t:
            continue
        if any(m in t for m in markers):
            paras.insert(0, doc.paragraphs[j])
        else:
            break
    return paras


def next_is_heading_or_empty(para) -> bool:
    """Чи наступний абзац — інший підрозділ без тексту між ними."""
    parent = para._element.getparent()
    nxt = para._element.getnext()
    if nxt is None:
        return True
    from docx.text.paragraph import Paragraph

    np = Paragraph(nxt, para._parent)
    t = np.text.strip()
    if not t:
        return True
    return bool(re.match(r"^\d+\.\d+(\.\d+)?\.\s", t))


def remove_orphan_empty_headings(doc) -> None:
    """Видалити підряд порожні заголовки 3.1.x перед основним текстом розділу 3."""
    to_del = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if re.match(r"^3\.1\.[1-6]\.", t) and p.style.name.startswith("Heading"):
            nxt = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None
            if nxt and re.match(r"^3\.1\.[1-6]\.", nxt.text.strip()):
                to_del.append(p)
            elif nxt and re.match(r"^3\.1\. ", nxt.text.strip()):
                to_del.append(p)
    for p in to_del:
        delete_paragraph(p)


def ensure_heading_before(doc, needle: str, heading: str, level: int = 3) -> None:
    for p in doc.paragraphs:
        if needle in p.text and not p.style.name.startswith("toc"):
            prev = None
            parent = p._element.getparent()
            prev_el = p._element.getprevious()
            if prev_el is not None:
                from docx.text.paragraph import Paragraph

                prev = Paragraph(prev_el, p._parent)
            if prev and heading[:10] in prev.text:
                return
            hp = p.insert_paragraph_before(heading)
            set_heading(hp, level)
            # прибрати дубль у першому реченні
            if p.text.strip().startswith(needle.split(".")[0] if "." in needle else needle):
                parts = p.text.split(".", 1)
                if len(parts) > 1 and len(parts[0]) < 80:
                    p.text = parts[1].strip()
            return


def fix_block(doc, heading_prefix: str, markers: tuple[str, ...]) -> None:
    h = find_heading(doc, heading_prefix)
    if not h:
        return
    body = collect_before(doc, h, markers)
    if not body:
        return
    texts = [p.text for p in body]
    for p in body:
        delete_paragraph(p)
    anchor = h
    for t in texts:
        anchor = add_body_after(anchor, t)


def expand_214(doc) -> None:
    h = find_heading(doc, "2.1.4. Користувацькі вимоги")
    if not h:
        return
    fix_block(
        doc,
        "2.1.4.",
        (
            "Користувацькі вимоги описують",
            "Диспетчерська служба (майбутній",
        ),
    )
    if any("UR-01" in p.text for p in doc.paragraphs):
        return
    texts = [
        "Користувацькі вимоги (User Requirements) формалізують очікування основних акторів підсистеми RoutePulse.",
        "Вимоги до інтерфейсу водія (мобільний термінал):",
    ]
    anchor = h
    for t in texts:
        anchor = add_body_after(anchor, t)
    bullets = [
        "UR-01: авторизація за DRV-NNNN і паролем не більше ніж за 3 дії на екрані;",
        "UR-02: старт рейсу (вибір маршруту та ТЗ) — не більше 4 кроків;",
        "UR-03: під час руху — мінімум ручного введення; статус рейсу (in_progress/completed) зрозумілий без додаткових пояснень;",
        "UR-04: реєстрація інциденту з фото та геоприв'язкою — не більше 5 кроків;",
        "UR-05: при втраті мережі GPS-дані не втрачаються (офлайн-черга Room + синхронізація);",
        "UR-06: розмір шрифту ≥16 sp, кнопки ≥64 dp (NFR-03).",
    ]
    for b in bullets:
        anchor = add_bullet_after(anchor, b)
    anchor = add_body_after(
        anchor,
        "Вимоги диспетчерської служби (майбутній веб-клієнт): перегляд live-карти активних рейсів, "
        "журналу інцидентів, архіву waybill і експорт звітів — через REST API без дублювання введення даних водієм.",
    )
    # прибрати старі короткі дублікати перед 2.1.3
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("Користувацькі вимоги описують очікування водія") and "UR-01" not in t:
            delete_paragraph(p)
        if t.startswith("Диспетчерська служба (майбутній веб") and "перегляд live" not in t:
            delete_paragraph(p)


def main() -> None:
    doc = Document(str(THESIS))

    fix_block(
        doc,
        "1.3.",
        (
            "На основі проведеного аналізу предметної галузі",
            "Клієнтська частина реалізована",
            "MongoDB 6.x обрана",
            "Моделювання виконано за UML",
            "обрані технології повністю відповідають",
            "Серверна частина побудована",
            "Для офлайн-режиму обрано Room",
        ),
    )

    fix_block(
        doc,
        "2.2.3.",
        (
            "Алгоритм реєстрації інциденту не передбачає",
            "Умова 1 (валідація опису)",
            "Умова 2 (вибір зупинки)",
            "Умова 3 (успішна",
            "формула гаверсинуса",
            "Батч телеметрії",
            "d = 2R",
        ),
    )
    h223 = find_heading(doc, "2.2.3.")
    if h223:
        intro = (
            "Математичне забезпечення підсистеми включає формалізовані умови валідації вхідних даних "
            "при реєстрації інциденту та геометричні розрахунки для обробки GPS-телеметрії."
        )
        nxt = h223._element.getnext()
        if nxt is not None:
            from docx.text.paragraph import Paragraph

            first = Paragraph(nxt, h223._parent).text.strip()
            if intro[:40] not in first:
                add_body_after(h223, intro)

    fix_block(
        doc,
        "2.1.4.",
        (
            "Користувацькі вимоги описують",
            "Диспетчерська служба (майбутній",
        ),
    )
    expand_214(doc)

    remove_orphan_empty_headings(doc)

    # прибрати дубль 3.1.3/таблицю з початку розділу 3 (перед 3.1. Інформаційне)
    h31 = find_heading(doc, "3.1. Інформаційне")
    if h31:
        to_remove = []
        el = h31._element.getprevious()
        while el is not None:
            from docx.text.paragraph import Paragraph

            pp = Paragraph(el, h31._parent)
            t = pp.text.strip()
            if re.match(r"^3\.1\.[2-6]\.", t) or "Таблиця 3.1" in t or "Джерела вхідних" in t:
                to_remove.append(pp)
                el = el.getprevious()
            else:
                break
        for p in to_remove:
            delete_paragraph(p)

    ensure_heading_before(
        doc,
        "Загальна характеристика інформаційного забезпечення",
        "3.1.1. Загальна характеристика інформаційного забезпечення",
    )
    ensure_heading_before(
        doc,
        "Побудова системи класифікації та кодування",
        "3.1.2. Побудова системи класифікації та кодування",
    )
    ensure_heading_before(
        doc,
        "Організація збору і передачі первинної інформації",
        "3.1.3. Джерела вхідних даних та форми їх подання",
    )
    ensure_heading_before(
        doc,
        "Інфологічна модель бази даних",
        "3.1.4. Інфологічна модель бази даних",
    )
    ensure_heading_before(
        doc,
        "Даталогічна модель бази даних",
        "3.1.5. Даталогічна модель бази даних",
    )
    ensure_heading_before(
        doc,
        "Структура інформаційних масивів",
        "3.1.6. Структура інформаційних масивів",
    )

    # таблиця джерел даних — під 3.1.3
    for p in doc.paragraphs:
        if "Джерела вхідних даних та форми їх подання наведено в таблиці 3.1" in p.text:
            prev_el = p._element.getprevious()
            if prev_el is not None:
                from docx.text.paragraph import Paragraph

                prev = Paragraph(prev_el, p._parent)
                if "3.1.2." in prev.text:
                    move_elements_after(p, [prev])
            break

    out = OUT
    try:
        doc.save(str(THESIS))
        out = THESIS
    except PermissionError:
        doc.save(str(OUT))
        print(f"Word відкрив файл — збережено як: {OUT}")
    else:
        print(f"Оновлено: {THESIS}")
    print("Sections filled and reordered.")


if __name__ == "__main__":
    main()
