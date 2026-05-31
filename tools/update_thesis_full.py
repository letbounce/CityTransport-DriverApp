# -*- coding: utf-8 -*-
"""Повне оновлення Коксюк_Диплом.docx: анотація, таблиці, 3.5, додатки, графіка."""
from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

THESIS = Path(r"c:\Users\letbounce\Desktop\Коксюк_Диплом.docx")
BACKUP = Path(r"c:\Users\letbounce\Desktop\Коксюк_Диплом_backup.docx")
SRC = Path(r"c:\Users\letbounce\Desktop\Дипломка")
FIG_W = Cm(14)


def style_table(table) -> None:
    try:
        table.style = "Table Grid"
    except KeyError:
        pass


# MongoDB incident screenshot (typo in filename on Desktop)
INCIDENT_DB_FILE = "incudents в Mongo DB.png"


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_heading_style(paragraph: Paragraph, level: int) -> None:
    paragraph.style = f"Heading {level}"


def add_body(doc: Document, text: str, indent: bool = True) -> Paragraph:
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p


def add_caption(doc: Document, text: str) -> Paragraph:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(12)
    return p


def add_figure(doc: Document, filename: str, caption: str, explanation: str | None = None) -> None:
    path = SRC / filename
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=FIG_W)
    else:
        doc.add_paragraph(f"[Файл не знайдено: {filename}]")
    add_caption(doc, caption)
    if explanation:
        add_body(doc, explanation)
    doc.add_paragraph()


def add_table_with_caption(
    doc: Document,
    caption: str,
    headers: list[str],
    rows: list[list[str]],
) -> None:
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    style_table(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def insert_annotation(doc: Document) -> None:
    if any("АНОТАЦІЯ" in p.text.upper() for p in doc.paragraphs[:30]):
        return
    ref = next((p for p in doc.paragraphs if p.text.strip() == "РЕФЕРАТ"), None)
    if not ref:
        return

    blocks = [
        ("АНОТАЦІЯ", 1),
        (
            "випускного бакалаврського проекту студента 4 курсу\n"
            "Навчально-наукового інституту «Інститут інформаційних технологій в економіці»\n"
            "Коксюка Олега Віталійовича, виконаного на тему:\n"
            "«Інформаційна система диспетчеризації міського пасажирського транспорту "
            "з використанням мобільних платформ»\n"
            "Київ: кафедра інформаційних систем в економіці, 2026 р.",
            0,
        ),
        (
            "Випускний бакалаврський проект присвячено автоматизації роботи водія "
            "автотранспортного підприємства: електронний дорожній лист, GPS-трекінг, "
            "реєстрація інцидентів із фотофіксацією та передача телеметрії на сервер. "
            "Розроблено прототип підсистеми RoutePulse (Android + Node.js/Express + MongoDB).",
            0,
        ),
        (
            "Робота складається з трьох розділів. У розділі 1 проаналізовано предметну галузь "
            "та існуючі ІС диспетчеризації. У розділі 2 сформульовано вимоги (FR/NFR), "
            "постановку задачі, алгоритм обробки інциденту та UML-моделі. У розділі 3 "
            "наведено інформаційне, технічне, програмне та організаційне забезпечення, "
            "результати апробації на контрольному прикладі DRV-1042.",
            0,
        ),
        (
            "Новизна роботи полягає в поєднанні офлайн-first архітектури мобільного клієнта "
            "(Room + WorkManager), бакетного зберігання телеметрії в MongoDB та версійного "
            "обліку змін інцидентів у self-hosted рішенні для АТП без ліцензійних платежів.",
            0,
        ),
        (
            "Практична значущість: прототип може бути впроваджений в АТП після доопрацювання "
            "веб-кабінету диспетчера; підтверджено smoke-тестуванням 14 API-ендпоінтів "
            "(див. рис. Г.12–Г.13, табл. 3.5).",
            0,
        ),
        (
            "Ключові слова: інформаційна система, диспетчеризація, мобільний термінал водія, "
            "Android, MongoDB, GPS-трекінг, дорожній лист, інцидент, REST API, RoutePulse.",
            0,
        ),
    ]
    for text, level in reversed(blocks):
        if level:
            p = ref.insert_paragraph_before(text)
            p.style = f"Heading {level}"
        else:
            p = ref.insert_paragraph_before("")
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(text)
    ref.insert_paragraph_before("")


def fix_section3_headings(doc: Document) -> None:
    patterns = [
        (r"^3\.2\. Технічне", 2),
        (r"^3\.2\.[1-4]\.", 3),
        (r"^3\.3\. Програмне", 2),
        (r"^3\.3\.[1-4]\.", 3),
        (r"^3\.4\. Результати", 2),
        (r"^3\.4\.[1-3]\.", 3),
        (r"^3\.5\. Організаційне", 2),
    ]
    for p in doc.paragraphs:
        t = p.text.strip()
        for pat, lvl in patterns:
            if re.match(pat, t):
                set_heading_style(p, lvl)
                break


def insert_after_paragraph(doc: Document, needle: str, new_paragraphs: list[str]) -> bool:
    for p in doc.paragraphs:
        if needle in p.text:
            anchor = p
            for text in reversed(new_paragraphs):
                np = anchor.insert_paragraph_before(text)
                np.paragraph_format.first_line_indent = Cm(1.25)
                np.paragraph_format.line_spacing = 1.5
                np.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            return True
    return False


def add_section_225_tables(doc: Document) -> None:
    needle = "Перелік масивів використовуваної інформації наведено в таблиці 2.4"
    for p in doc.paragraphs:
        if needle in p.text:
            anchor = p
            inserts = [
                "Використовувана інформація. Для алгоритму реєстрації інциденту використовуються "
                "масиви згідно з таблицею 2.4.",
                "",
                "Таблиця 2.4",
                "Перелік масивів використовуваної інформації (алгоритм реєстрації інциденту)",
                "",
                "Результати розв'язання. Після успішного виконання алгоритму формуються масиви "
                "результатної інформації згідно з таблицею 2.5.",
                "",
                "Таблиця 2.5",
                "Перелік масивів результатної інформації (алгоритм реєстрації інциденту)",
            ]
            for t in reversed(inserts):
                np = anchor.insert_paragraph_before(t)
                if t.startswith("Таблиця"):
                    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif t:
                    np.paragraph_format.first_line_indent = Cm(1.25)
                    np.paragraph_format.line_spacing = 1.5
                    np.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            break

    # Add table 2.5 after table 2.4 if not exists
    if not any("Таблиця 2.5" in p.text and "результат" in p.text.lower() for p in doc.paragraphs):
        # find "Математичний опис"
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip().startswith("Математичний опис"):
                doc.paragraphs[i].insert_paragraph_before("")
                # add table before math - search backwards for table 2.4 content end
                break

    # Simpler: append table 2.5 via add at math description
    for p in doc.paragraphs:
        if p.text.strip().startswith("Математичний опис"):
            before = p.insert_paragraph_before("")
            rows = [
                ["Документ інциденту в MongoDB", "INCIDENT_DOC", "1"],
                ["URL збереженого фото", "PHOTO_URL", "1"],
                ["Підтвердження UI (екран рейсу)", "UI_ACK", "1"],
                ["HTTP-відповідь сервера", "HTTP_201", "1"],
            ]
            # insert table XML before paragraph - use add_table at end then move is hard
            # use insert_paragraph_before for caption only; add_table after doc built
            cap = p.insert_paragraph_before("Таблиця 2.5\nПерелік масивів результатної інформації (алгоритм реєстрації інциденту)")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tbl = doc.add_table(rows=1 + len(rows), cols=3)
            style_table(tbl)
            hdr = ["Масив", "Ідентифікатор", "Макс. к-сть записів"]
            for j, h in enumerate(hdr):
                tbl.rows[0].cells[j].text = h
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row):
                    tbl.rows[ri + 1].cells[ci].text = val
            p._element.addprevious(tbl._tbl)
            p.insert_paragraph_before("")
            break


def add_traceability_matrix(doc: Document) -> None:
    if any("Таблиця 2.6" in p.text for p in doc.paragraphs):
        return
    for p in doc.paragraphs:
        if "2.3.3." in p.text and "Розподіл вимог" in p.text:
            continue
    anchor = None
    for p in doc.paragraphs:
        if "FR-01 (Авторизація)" in p.text or "трасування підтверджує" in p.text:
            anchor = p
    if not anchor:
        for p in doc.paragraphs:
            if "Рисунок 2.10" in p.text and "трасування" in p.text.lower():
                anchor = p
    if not anchor:
        return

    texts = [
        "Додатково до діаграми трасування (рис. 2.10) наведено матрицю відповідності вимог, "
        "прецедентів і компонентів реалізації (табл. 2.6).",
        "Таблиця 2.6",
        "Матриця трасування вимог до компонентів підсистеми RoutePulse",
    ]
    for t in reversed(texts):
        np = anchor.insert_paragraph_before(t)
        if t.startswith("Таблиця"):
            np.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            np.paragraph_format.first_line_indent = Cm(1.25)
            np.paragraph_format.line_spacing = 1.5
            np.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    rows = [
        ["FR-01", "UC-01", "LoginScreen, AuthController"],
        ["FR-02", "UC-02", "RouteDashboardScreen, RouteController"],
        ["FR-03", "UC-02", "WaybillController, ActiveTripScreen"],
        ["FR-04", "UC-06", "ActiveTripScreen, PATCH /complete"],
        ["FR-05", "UC-06", "Archive API, WaybillController"],
        ["FR-06", "UC-05", "IncidentReportScreen, IncidentController"],
        ["FR-07", "UC-05", "IncidentEditScreen, version_history"],
        ["FR-08", "UC-03", "TripMapRouteScreen, OSMdroid"],
        ["FR-09", "UC-04", "LocationTrackingService, TelemetryController"],
        ["FR-10", "UC-03", "GET /api/map/live-trips"],
        ["NFR-01", "—", "Room LocalQueue, TelemetrySyncWorker"],
        ["NFR-02", "—", "JWT, EncryptedSharedPreferences"],
        ["NFR-03", "—", "Compose UI (≥16 sp, ≥64 dp)"],
        ["NFR-04", "—", "GPS 5 с, batch 100 точок"],
    ]
    tbl = doc.add_table(rows=1 + len(rows), cols=3)
    style_table(tbl)
    for j, h in enumerate(["Вимога", "Прецедент", "Компонент реалізації"]):
        tbl.rows[0].cells[j].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            tbl.rows[ri + 1].cells[ci].text = val
    anchor._element.addprevious(tbl._tbl)
    anchor.insert_paragraph_before("")


def add_table_31_waybills(doc: Document) -> None:
    if any("Таблиця 3.1" in p.text and "waybills" in p.text.lower() for p in doc.paragraphs):
        return
    if any("Таблиця 3.1" in p.text for p in doc.paragraphs):
        return
    anchor = None
    for p in doc.paragraphs:
        if "Структура інформаційних масивів" in p.text:
            anchor = p
            break
    if not anchor:
        return
    intro = (
        "Опис структури масиву waybills (колекція MongoDB) наведено в таблиці 3.1 "
        "відповідно до вимог РД 50-34.698-90."
    )
    np = anchor.insert_paragraph_before(intro)
    np.paragraph_format.first_line_indent = Cm(1.25)
    np.paragraph_format.line_spacing = 1.5
    np.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    cap = anchor.insert_paragraph_before("Таблиця 3.1\nОпис структури масиву waybills")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rows = [
        ["driver_id", "String", "PK", "так", "ні", "Ідентифікатор водія DRV-NNNN"],
        ["route_id", "ObjectId", "FK", "так", "ні", "Посилання на routes"],
        ["route_number", "String", "—", "так", "ні", "Номер маршруту"],
        ["status", "enum", "—", "так", "ні", "assigned|in_progress|completed|cancelled"],
        ["vehicle_id", "String", "—", "так", "ні", "Код ТЗ, напр. KP-4412"],
        ["started_at", "Date", "—", "так", "ні", "Час початку рейсу"],
        ["completed_at", "Date", "—", "ні", "так", "Час завершення рейсу"],
    ]
    headers = ["Поле", "Тип", "Ключ", "Обов'язкове", "Індекс", "Примітка"]
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    style_table(tbl)
    for j, h in enumerate(headers):
        tbl.rows[0].cells[j].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            tbl.rows[ri + 1].cells[ci].text = val
    anchor._element.addprevious(tbl._tbl)
    anchor.insert_paragraph_before("")


def add_section_35(doc: Document) -> None:
    if any("3.5." in p.text and "Організаційне" in p.text for p in doc.paragraphs):
        return
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip() == "ВИСНОВКИ":
            anchor = p
            break
    if not anchor:
        return

    parts = [
        ("3.5. Організаційне забезпечення", 2),
        (
            "Організаційне забезпечення визначає розподіл ролей, регламенти взаємодії "
            "та умови експлуатації підсистеми RoutePulse в АТП.",
            0,
        ),
        (
            "Ролі користувачів: водій (~150 осіб) — первинне введення даних через мобільний "
            "термінал; диспетчер зміни — моніторинг активних рейсів (через майбутній веб-кабінет); "
            "диспетчер з безпеки руху — обробка інцидентів; адміністратор ІС — seed, резервне "
            "копіювання MongoDB, оновлення APK.",
            0,
        ),
        (
            "Регламент роботи водія: (1) авторизація на початку зміни; (2) створення waybill "
            "перед виїздом (рис. Г.3–Г.4); (3) ведення активного рейсу з GPS (рис. Г.4); "
            "(4) завершення рейсу кнопкою «Завершити рейс»; (5) реєстрація інцидентів за потреби.",
            0,
        ),
        (
            "Навчання та супровід: інструктаж 2 год (UI, офлайн-режим, фотофіксація); "
            "технічна підтримка — адміністратор сервера; оновлення застосунку — централізована "
            "роздача APK або Google Play (внутрішній канал).",
            0,
        ),
        (
            "Таблиця 3.6",
            0,
        ),
        (
            "Організаційні заходи впровадження підсистеми RoutePulse",
            0,
        ),
    ]
    for text, level in reversed(parts):
        if level:
            p = anchor.insert_paragraph_before(text)
            p.style = f"Heading {level}"
        elif text == "Таблиця 3.6":
            p = anchor.insert_paragraph_before(text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif text.startswith("Організаційні"):
            p = anchor.insert_paragraph_before(text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p = anchor.insert_paragraph_before("")
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(text)

    rows = [
        ["Пілотне впровадження", "1 АТП, 5–10 ТЗ", "2 тижні", "Водії, диспетчер"],
        ["Навчання персоналу", "150 водіїв", "2 год/особа", "Інструктор ІС"],
        ["Резервне копіювання БД", "Щоденно", "mongodump", "Адміністратор"],
        ["Оновлення мобільного ПЗ", "Щомісяця", "APK", "Адміністратор"],
    ]
    tbl = doc.add_table(rows=1 + len(rows), cols=4)
    style_table(tbl)
    for j, h in enumerate(["Захід", "Охоплення", "Періодичність", "Відповідальний"]):
        tbl.rows[0].cells[j].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            tbl.rows[ri + 1].cells[ci].text = val
    anchor._element.addprevious(tbl._tbl)
    anchor.insert_paragraph_before("")


def update_text_references(doc: Document) -> None:
    replacements = [
        (
            "на рисунку Г.5 підтверджено збереження waybill",
            "на рисунках Г.12–Г.13 підтверджено збереження даних waybill та incident",
        ),
        (
            "Інтерфейс мобільного клієнта наведено на рисунках Г.1–Г.4",
            "Інтерфейс мобільного клієнта наведено на рисунках Г.1–Г.11",
        ),
        (
            "live-маркери активних рейсів",
            "трек маршруту та зупинки",
        ),
        (
            "на рис. 2.12",
            "на рис. 2.10",
        ),
    ]
    for p in doc.paragraphs:
        for old, new in replacements:
            if old in p.text:
                p.text = p.text.replace(old, new)


def remove_appendices_and_graphics(doc: Document) -> None:
    to_del = []
    started = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "ДОДАТКИ" and p.style.name.startswith("Heading"):
            started = True
        if started:
            to_del.append(p)
    for p in to_del:
        delete_paragraph(p)


GRAPHIC_ITEMS = [
    ("Вхід.png", "Рисунок Г.1 — Екран авторизації водія (LoginScreen)", "FR-01."),
    (
        "Головне меню.png",
        "Рисунок Г.2 — Головне меню (HomeMenuScreen)",
        "Навігація після авторизації.",
    ),
    (
        "Створення дорожнього листа.png",
        "Рисунок Г.3 — Створення дорожнього листа (RouteDashboardScreen)",
        "Вибір маршруту №114 та ТЗ. FR-02, FR-03.",
    ),
    (
        "Активний рейс.png",
        "Рисунок Г.4 — Активний рейс із кнопкою «Завершити рейс» (ActiveTripScreen)",
        "Відображення поточного waybill, GPS-трекінг. FR-03, FR-04, FR-09.",
    ),
    (
        "Лист рейсів.png",
        "Рисунок Г.5 — Перелік дорожніх листів водія",
        "Перегляд активних та завершених рейсів.",
    ),
    (
        "Редагування дорожнього листа.png",
        "Рисунок Г.6 — Редагування параметрів дорожнього листа",
        "PATCH /api/waybills/:id.",
    ),
    (
        "Мапа маршруту.png",
        "Рисунок Г.7 — Карта маршруту №114 (TripMapRouteScreen, OSMdroid)",
        "FR-08.",
    ),
    (
        "Створення звіту про інцидент.png",
        "Рисунок Г.8 — Реєстрація інциденту (IncidentReportScreen)",
        "FR-06.",
    ),
    (
        "Редагування інциденту.png",
        "Рисунок Г.9 — Редагування інциденту",
        "FR-07, version_history.",
    ),
    (
        "Архівування інциденту.png",
        "Рисунок Г.10 — Архівування інциденту",
        "soft-delete, deleted_at.",
    ),
    (
        "Архів інцидентів.png",
        "Рисунок Г.11 — Архів інцидентів (IncidentsListScreen)",
        "",
    ),
    (
        "waybill_completed.png",
        "Рисунок Г.12 — Документ waybills у MongoDB (статус completed, DRV-1042)",
        "Контрольний приклад, п. 3.4.",
    ),
    (
        INCIDENT_DB_FILE,
        "Рисунок Г.13 — Документ incidents у MongoDB після реєстрації інциденту",
        "Підтвердження FR-06 на контрольному прикладі.",
    ),
]

APPENDIX_CODE = [
    ("Waybill.png", "Рисунок А.1 — Схема Mongoose Waybill.js"),
    ("Incident.png", "Рисунок А.2 — Схема Mongoose Incident.js"),
]

REQUIREMENTS_ROWS = [
    ["FR-01", "Авторизація", "Обов'язкова", "JWT 8 год, DRV-NNNN", "UC-01", "LoginScreen"],
    ["FR-02", "Маршрути", "Обов'язкова", "Список маршрутів і зупинок", "UC-02", "RouteController"],
    ["FR-03", "Створення waybill", "Обов'язкова", "status=in_progress", "UC-02", "WaybillController"],
    ["FR-04", "Завершення рейсу", "Обов'язкова", "completed_at, stop GPS", "UC-06", "ActiveTripScreen"],
    ["FR-05", "Архівування", "Бажана", "5 кодів причини", "UC-06", "Archive API"],
    ["FR-06", "Інцидент+фото", "Обов'язкова", "Base64, GPS", "UC-05", "IncidentController"],
    ["FR-07", "Редагування", "Бажана", "version_history", "UC-05", "IncidentEditScreen"],
    ["FR-08", "Карта", "Обов'язкова", "OSM, GeoJSON", "UC-03", "TripMapRouteScreen"],
    ["FR-09", "GPS-трекінг", "Обов'язкова", "кожні 5 с", "UC-04", "LocationTrackingService"],
    ["FR-10", "Live API", "Бажана", "GET /map/live-trips", "UC-03", "Map API"],
    ["NFR-01", "Офлайн", "Обов'язкова", "Room+Worker, batch 100", "—", "TelemetrySyncWorker"],
    ["NFR-02", "Безпека", "Обов'язкова", "JWT, bcrypt", "—", "AuthController"],
    ["NFR-03", "UX", "Обов'язкова", "≥16 sp, ≥64 dp", "—", "Compose UI"],
    ["NFR-04", "Продуктивність", "Обов'язкова", "GPS 5 с, sync 15 хв", "—", "FGS+Worker"],
]

CONTROL_DATA = [
    ["drivers", "3", "DRV-1042 (тестовий)"],
    ["routes", "3", "№114 (8 зупинок)"],
    ["vehicles", "5", "KP-4412, електробус"],
    ["waybills", "1+", "completed для DRV-1042"],
    ["incidents", "1+", "з GPS та photo_url"],
    ["telemetry", "бакети", "~100 точок/документ"],
]


def build_appendices_and_graphics(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("ДОДАТКИ", level=1)
    add_body(
        doc,
        "У додатках наведено фрагменти реалізації, деталізовану специфікацію вимог, "
        "опис структури даних та контрольний приклад наповнення БД routepulse_db.",
    )

    doc.add_heading("Додаток А", level=2)
    doc.add_heading("Фрагменти моделей даних (Mongoose)", level=3)
    for fname, cap in APPENDIX_CODE:
        add_figure(doc, fname, cap)

    doc.add_heading("Додаток Б", level=2)
    doc.add_heading("Приклад PDF-звіту про інцидент", level=3)
    add_figure(
        doc,
        "Приклад звіту інциденту в PDF.png",
        "Рисунок Б.1 — Приклад сформованого PDF-звіту про інцидент",
        "Вихідний документ для служби безпеки руху.",
    )

    doc.add_heading("Додаток В", level=2)
    doc.add_heading("Деталізована специфікація вимог (ISO/IEC/IEEE 29148:2018)", level=3)
    add_table_with_caption(
        doc,
        "Таблиця В.1\nПовна специфікація функціональних та нефункціональних вимог",
        ["ID", "Назва", "Пріоритет", "Критерій прийняття", "UC", "Компонент"],
        REQUIREMENTS_ROWS,
    )

    doc.add_heading("Додаток Г", level=2)
    doc.add_heading("Структура даних та контрольний приклад", level=3)
    add_body(
        doc,
        "Фрагмент логічного опису колекції waybills (аналог DDL для документної СКБД MongoDB): "
        "driver_id (String, PK), route_id (ObjectId, FK→routes), route_number (String), "
        "status (enum: assigned|in_progress|completed|cancelled), vehicle_id (String), "
        "started_at/completed_at (Date), deleted_at, deletion_reason_code (soft-delete).",
    )
    add_body(
        doc,
        "Колекція incidents: waybill_id (FK), driver_id, type (accident|breakdown|traffic_jam|other), "
        "description, location{lat,lng}, photo_url, reported_at, stop_label, "
        "can_move_independently, status, version_history[], is_modified.",
    )
    add_table_with_caption(
        doc,
        "Таблиця Г.1\nКонтрольний приклад наповнення БД (npm run seed)",
        ["Колекція", "К-сть записів", "Приклад"],
        CONTROL_DATA,
    )

    doc.add_page_break()
    doc.add_heading("ГРАФІЧНІ МАТЕРІАЛИ", level=1)
    add_body(
        doc,
        "Наведено екранні форми мобільного застосунку RoutePulse та підтвердження збереження "
        "даних у MongoDB за контрольним прикладом (DRV-1042, маршрут №114). "
        "Ілюстрації Г.1–Г.11 — інтерфейс; Г.12–Г.13 — апробація в БД.",
    )
    doc.add_heading("Інтерфейс та результати апробації", level=2)
    for fname, cap, expl in GRAPHIC_ITEMS:
        add_figure(doc, fname, cap, expl or None)


def patch_apribation_paragraph(doc: Document) -> None:
    for p in doc.paragraphs:
        if "3.4.2." in p.text or "Апробація проводилась" in p.text:
            if "рис. Г.12" not in p.text and "Апробація" in p.text:
                p.text = (
                    p.text.rstrip()
                    + " Підтвердження збереження даних у MongoDB наведено на рисунках Г.12 (waybill) "
                    "та Г.13 (incident); екран активного рейсу з завершенням — на рисунку Г.4."
                )
            break
        if "Усі 14 API" in p.text:
            p.text = (
                p.text
                + " Інтерфейс апробації ілюстровано рисунками Г.1–Г.11; документи БД — рис. Г.12–Г.13."
            )
            break


def main() -> None:
    if not THESIS.exists():
        raise SystemExit(f"Не знайдено: {THESIS}")
    shutil.copy2(THESIS, BACKUP)
    print(f"Резервна копія: {BACKUP}")

    doc = Document(str(THESIS))
    insert_annotation(doc)
    add_section_225_tables(doc)
    add_traceability_matrix(doc)
    add_table_31_waybills(doc)
    add_section_35(doc)
    update_text_references(doc)
    patch_apribation_paragraph(doc)
    fix_section3_headings(doc)
    remove_appendices_and_graphics(doc)
    build_appendices_and_graphics(doc)
    doc.save(str(THESIS))
    print(f"Оновлено: {THESIS}")


if __name__ == "__main__":
    main()
