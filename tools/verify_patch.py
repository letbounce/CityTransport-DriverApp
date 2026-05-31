# -*- coding: utf-8 -*-
from docx import Document
import re

doc = Document(r"d:\Downloads\Коксюк_Диплом.docx")
checks = {
    "ANOTACIYA body": sum(1 for p in doc.paragraphs if p.text.strip() == "АНОТАЦІЯ" and not p.style.name.startswith("toc")),
    "REFERAT": any(p.text.strip() == "РЕФЕРАТ" for p in doc.paragraphs),
    "3.1.6": any("3.1.6" in p.text for p in doc.paragraphs if not p.style.name.startswith("toc")),
    "3.5": any(re.match(r"^3\.5\.", p.text.strip()) for p in doc.paragraphs if not p.style.name.startswith("toc")),
    "3.4.4": any("3.4.4" in p.text for p in doc.paragraphs),
    "galuz (predmet)": sum(1 for p in doc.paragraphs if "предметн" in p.text.lower() and "галуз" in p.text.lower()),
    "oblast": sum(1 for p in doc.paragraphs if "предметн" in p.text.lower() and "област" in p.text.lower()),
    "intro theory": any("Теоретична" in p.text and "методична" in p.text for p in doc.paragraphs),
    "bib 32": 0,
    "citations [N]": sum(1 for p in doc.paragraphs if re.search(r"\[\d+\]", p.text)),
}
after = False
n = 0
for p in doc.paragraphs:
    if p.text.strip() == "ПЕРЕЛІК ВИКОРИСТАНИХ ДЖЕРЕЛ":
        after = True
        continue
    if after:
        if p.text.strip().startswith("ДОДАТКИ"):
            break
        if re.match(r"^\d+\.\s", p.text.strip()):
            n += 1
checks["bib 32"] = n
# table V.1 cols
for t in doc.tables:
    if t.rows and t.rows[0].cells[0].text.strip() == "ID" and len(t.rows) > 1 and t.rows[1].cells[0].text == "FR-01":
        checks["table V cols"] = len(t.columns)
        checks["table V headers"] = [c.text for c in t.rows[0].cells]
        break
for k, v in checks.items():
    print(f"{k}: {v}")
